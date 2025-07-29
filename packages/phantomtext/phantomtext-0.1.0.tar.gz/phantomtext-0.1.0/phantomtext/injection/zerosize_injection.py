from ..attack_base_injection import AttackBase
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader, PdfWriter
from io import BytesIO
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup

class ZeroSizeInjection(AttackBase):
    def __init__(self, modality="default", file_format="pdf"):
        if file_format == "pdf" or file_format == "html":
            self.file_format = file_format

            if modality == "default":
                self.modality = modality
                self.font_size=0
            
            elif modality == 'close-to-zero':
                self.modality = modality
                self.font_size=0.1

            # Modality not valid, raise error
            else:
                raise ValueError(f"Modality {modality} is not valid. Valid modality is 'default' (font-size = 0 ) or 'close-to-zero' (font-size=0.1)")
        elif file_format == "docx":
            self.file_format = file_format
            if modality == "default":
                self.modality = modality
                self.font_size=1
            else:
                raise ValueError(f"Modality {modality} is not valid. Valid modality is 'default' for docx (font-size = 1, minimum possible )")
        else:
            raise ValueError(f"File format {file_format} is not valid. Valid file format is 'pdf', 'docx' or 'html'")
        # super.__init__(modality, file_format)

    def apply(self, input_document, injection, font_size=12, x_coord=100, y_coord=730, image_file=None, output_path=None):
        # Check if input document exists
        if not input_document:
            raise ValueError("Input document is required")
        
        # Check if injection exists
        if not injection:
            raise ValueError("Injection is required")
        
        # Check if x and y coordinates are valid
        if x_coord < 0 or y_coord < 0:
            raise ValueError("x and y coordinates must be positive")
        
        # Check if output path is provided
        if not output_path:
            output_path = f"{input_document.split('.')[0]}_injected.{self.file_format}"
        
        if self.file_format == "pdf":
            
            buffer = BytesIO()
            c = canvas.Canvas(buffer)

            # Write hidden text (white text matching background)
            c.setFontSize(self.font_size)
            c.drawString(x_coord, y_coord, injection)  # Adjust coordinates as needed

            c.save()
            buffer.seek(0)

            # Read the existing PDF
            existing_pdf = PdfReader(input_document)
            new_pdf = PdfReader(buffer)

            # Create a PdfWriter object to combine PDFs
            pdf_writer = PdfWriter()

            # Merge the new content with the existing PDF
            if len(existing_pdf.pages) > 0:
                existing_page = existing_pdf.pages[0]  # Assuming there's only one page
                new_page = new_pdf.pages[0]  # We only have one page from ReportLab
                
                existing_page.merge_page(new_page)  # Merge new content onto existing page
                
                pdf_writer.add_page(existing_page)

            # Write out the combined PDF to a file
            with open(output_path, 'wb') as f:
                pdf_writer.write(f)

        elif self.file_format == "docx":
            doc = Document(input_document)

            # Add hidden text (zero-size text)
            hidden_paragraph = doc.add_paragraph()
            hidden_run = hidden_paragraph.add_run(injection)
            hidden_run.font.size = Pt(self.font_size)  # Set font size to zero

            # Save the updated DOCX
            doc.save(output_path)

        elif self.file_format == "html":
            soup = BeautifulSoup(open(input_document), 'html.parser')
    
            # Create a new tag (e.g., <div>, <p>, <span>)
            new_tag = soup.new_tag('div')
            
            # Set the text content)
            new_tag.string = injection
            
            # Apply CSS to set font-size to 0 (invisible)
            new_tag['style'] = f"font-size: {self.font_size};"
            
            soup.body.append(new_tag)
            
            with open(output_path, 'w') as f:
                f.write(str(soup))
        
    def check(self, input_document):
        pass