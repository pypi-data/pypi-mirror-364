import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn # To set complex script fonts like Arabic/East Asian


class DOCXHandler:
    def read_docx(self, file_path):
        """
        Reads text from a DOCX file.

        Args:
            file_path (str): Path to the DOCX file.

        Returns:
            str: Extracted text from the DOCX file.
        """
        text = ""
        try:
            document = Document(file_path)
            for paragraph in document.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")
        return text

    def write_docx(self, file_path, content, font_name='DejaVuSans'):
        """
        Saves the given text as a DOCX file.

        Args:
            text (str): The text to save.
            file_path (str): The path to save the DOCX file.
        """
        try:
            # --- 1. Create Document ---
            document = Document()

            # --- 2. Add Paragraph with Unicode Text ---
            paragraph = document.add_paragraph()

            # Add the text as a 'run'. A paragraph can have multiple runs with different formatting.
            run = paragraph.add_run(content)

            # --- 3. Attempt to Set Font ---
            # Note: This only works if the font 'font_name' is installed on the
            # system where the DOCX is opened.
            run.font.name = font_name

            # For broader compatibility, especially with complex scripts (like Arabic, CJK),
            # it's good practice to set the 'complex script' font hint too.
            # This requires understanding the underlying XML structure (using oxml).
            r = run._r # Get the underlying XML element for the run
            r.rPr.rFonts.set(qn('w:cs'), font_name) # Set complex script font
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name) # Set East Asian font hint

            # Optional: Set font size
            run.font.size = Pt(12)

            # --- 4. Save the Document ---
            document.save(file_path)

        except Exception as e:
            print(f"An error occurred: {e}")
