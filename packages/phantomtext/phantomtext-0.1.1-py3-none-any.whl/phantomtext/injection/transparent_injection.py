from ..attack_base_injection import AttackBase
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader, PdfWriter
from io import BytesIO
from docx import Document
from docx.shared import Pt,RGBColor
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement

class TransparentInjection(AttackBase):
    def __init__(self, modality="default", file_format="pdf"):
        
        if file_format == "pdf" or file_format == "html":
            self.file_format = file_format
            
            #default = background color
            if modality == "default":
                self.modality = modality
                
            elif modality == 'opacity-0':
                self.modality = modality

            elif modality == 'opacity-close-to-zero':
                self.modality = modality
                
            # Modality not valid, raise error
            else:
                raise ValueError(f"Modality {modality} is not valid. Valid modality is 'default' (background color injection ), 'opacity-0' (opacity set to 0) or 'opacity-close-to-zero' (opacity=0.1)")
            
        elif file_format == "docx":
            self.file_format = file_format
            if modality == "default":
                self.modality = modality
            elif modality == 'vanish':
                self.modality = modality
            
            else:
                raise ValueError(f"Modality {modality} is not valid. Valid modality is 'default' for docx (background color injection ) or 'vanish' (opacity set to 0)")
        
        else:
            raise ValueError(f"File format {file_format} is not valid. Valid file format is 'pdf', 'docx' or 'html'")
        

    def apply(self, input_document, injection, font_size=12, x_coord=100, y_coord=730, image_file=None, output_path=None):
        if not input_document:
            raise ValueError("Input document is required")
        
        # Check if injection exists
        if not injection:
            raise ValueError("Injection is required")
        
        # Check if x and y coordinates are valid
        if x_coord < 0 or y_coord < 0:
            raise ValueError("x and y coordinates must be positive")
        
        # Check if output path is provided
        if not output_path:
            output_path = f"{input_document.split('.')[0]}_injected.{self.file_format}"
        
        if self.file_format == "pdf":
            
            buffer = BytesIO()
            c = canvas.Canvas(buffer)

            # Write hidden text (white text matching background)
            if self.modality == 'default':
                c.setFillColorRGB(1, 1, 1) 
            
            elif self.modality == 'opacity-0':
                c.setFillAlpha(0)
            
            elif self.modality == 'opacity-close-to-zero':
                c.setFillAlpha(0.1)
            
    
            c.drawString(x_coord, y_coord, injection)  # Adjust coordinates as needed

            c.save()
            buffer.seek(0)

            # Read the existing PDF
            existing_pdf = PdfReader(input_document)
            new_pdf = PdfReader(buffer)

            # Create a PdfWriter object to combine PDFs
            pdf_writer = PdfWriter()

            # Merge the new content with the existing PDF
            if len(existing_pdf.pages) > 0:
                existing_page = existing_pdf.pages[0]  # Assuming there's only one page
                new_page = new_pdf.pages[0]  # We only have one page from ReportLab
                
                existing_page.merge_page(new_page)  # Merge new content onto existing page
                
                pdf_writer.add_page(existing_page)

            # Write out the combined PDF to a file
            with open(output_path, 'wb') as f:
                pdf_writer.write(f)

        elif self.file_format == "docx":
            doc = Document(input_document)

            if self.modality == 'default':
                hidden_paragraph = doc.add_paragraph()
                hidden_run = hidden_paragraph.add_run(injection)
                hidden_run.font.color.rgb = RGBColor(255, 255, 255)  # White color

            elif self.modality == 'vanish':
                p = doc.add_paragraph()
                run = p.add_run(injection)

                # Set the text to be hidden
                rPr = run._r.get_or_add_rPr()
                vanishing = OxmlElement('w:vanish')
                rPr.append(vanishing)
            # Save the updated DOCX
            doc.save(output_path)

        elif self.file_format == "html":
            soup = BeautifulSoup(open(input_document), 'html.parser')
    
            # Create a new tag (e.g., <div>, <p>, <span>)
            new_tag = soup.new_tag('div')
            
            # Set the text content)
            new_tag.string = injection
            
            if self.modality == 'default':
                background_color='transparent'
                new_tag['style'] = f"color: {background_color}; background-color: {background_color};"
            elif self.modality == 'opacity-0':
                new_tag['style'] = f"opacity: 0;"
            elif self.modality == 'opacity-close-to-zero':
                new_tag['style'] = f"opacity: 0.1;"
            
            # Append the new tag to the body

            soup.body.append(new_tag)
            
            with open(output_path, 'w') as f:
                f.write(str(soup))
        
    
    def check(self, input_document):
        return True