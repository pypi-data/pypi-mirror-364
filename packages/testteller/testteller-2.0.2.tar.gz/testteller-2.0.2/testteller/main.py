import asyncio
import logging
import os
from functools import wraps

import typer
from typing_extensions import Annotated

from .generator_agent.agent import TestTellerRagAgent
from .config import settings
from .core.constants import (
    DEFAULT_OUTPUT_FILE, DEFAULT_COLLECTION_NAME, SUPPORTED_LLM_PROVIDERS,
    DEFAULT_CHROMA_PERSIST_DIRECTORY, SUPPORTED_TEST_OUTPUT_FORMATS,
    DEFAULT_TEST_OUTPUT_FORMAT, DEFAULT_TEST_GENERATION_DIR, APP_SHORT_DESCRIPTION
)
from pathlib import Path
# Import config modules inside function to avoid circular imports
from .core.utils.helpers import setup_logging
from .core.utils.loader import with_spinner
from ._version import __version__
from .core.utils.exceptions import EmbeddingGenerationError

# Import automation command functionality
try:
    from testteller.automator_agent import automate_command
    HAS_AUTOMATION = True
except ImportError:
    HAS_AUTOMATION = False


setup_logging()
logger = logging.getLogger(__name__)


async def save_test_cases_with_format(content: str, output_file: str, output_format: str) -> tuple[str, str]:
    """Save test cases in the specified format (md, pdf, docx).
    
    Returns:
        tuple: (actual_output_file, actual_format) - the file path and format actually used
    """
    output_path = Path(output_file)
    
    # Ensure the parent directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    if output_format == "md":
        # Save as markdown (default behavior)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return str(output_path), "md"
    elif output_format == "pdf":
        # Convert markdown to PDF
        try:
            # Try importing required libraries for PDF generation
            try:
                import markdown
                from weasyprint import HTML, CSS
                from io import StringIO
            except ImportError:
                # Fallback: save as markdown with a warning
                logger.warning("PDF libraries not available, saving as markdown instead")
                print("⚠️  PDF conversion libraries not installed. Saving as markdown.")
                markdown_path = output_path.with_suffix('.md')
                with open(markdown_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                return str(markdown_path), "md"
            
            # Convert markdown to HTML
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
            
            # Add enhanced CSS styling for better PDF appearance with tables
            css_content = """
            @page {
                margin: 1in;
                size: A4;
            }
            body {
                font-family: Arial, sans-serif;
                line-height: 1.6;
                color: #333;
            }
            h1, h2, h3 { 
                color: #2c3e50; 
                page-break-after: avoid;
            }
            h2 {
                border-bottom: 2px solid #3498db;
                padding-bottom: 5px;
            }
            pre, code {
                background-color: #f4f4f4;
                padding: 10px;
                border-radius: 5px;
                font-family: 'Courier New', monospace;
                font-size: 12px;
                word-wrap: break-word;
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 15px 0;
                font-size: 11px;
                page-break-inside: avoid;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
                vertical-align: top;
            }
            th { 
                background-color: #3498db;
                color: white;
                font-weight: bold;
                text-align: center;
            }
            tr:nth-child(even) { background-color: #f8f9fa; }
            tr:hover { background-color: #e8f4f8; }
            .test-summary { 
                margin-bottom: 30px; 
            }
            .test-details {
                margin-top: 30px;
                page-break-before: always;
            }
            """
            
            # Generate PDF
            HTML(string=html_content).write_pdf(
                output_path,
                stylesheets=[CSS(string=css_content)]
            )
            return str(output_path), "pdf"
        except Exception as e:
            logger.error(f"Failed to generate PDF: {e}")
            print(f"⚠️  Failed to generate PDF: {e}")
            print("Saving as markdown instead...")
            # Fallback to markdown
            markdown_path = output_path.with_suffix('.md')
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return str(markdown_path), "md"
    elif output_format == "docx":
        # Convert to DOCX format
        try:
            try:
                from docx import Document
                from docx.shared import Inches
                import re
            except ImportError:
                # Fallback: save as markdown with a warning
                logger.warning("DOCX libraries not available, saving as markdown instead")
                print("⚠️  DOCX conversion libraries not installed. Saving as markdown.")
                markdown_path = output_path.with_suffix('.md')
                with open(markdown_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                return str(markdown_path), "md"
            
            # Create a new Document
            doc = Document()
            
            # Split content into lines for processing
            lines = content.split('\n')
            
            current_table = None
            table_headers = []
            in_table = False
            
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                if not line:
                    # Add paragraph break for empty lines
                    doc.add_paragraph()
                elif line.startswith('# '):
                    # Heading 1
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    # Heading 2
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    # Heading 3
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    # Heading 4
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    # Numbered list
                    p = doc.add_paragraph(line, style='List Number')
                elif line.startswith('```'):
                    # Code block (simple handling)
                    p = doc.add_paragraph(line, style='Intense Quote')
                elif line.startswith('|') and '|' in line[1:]:
                    # Table row detected
                    if not in_table:
                        # Check if next line is table separator
                        if i + 1 < len(lines) and lines[i + 1].strip().startswith('|---'):
                            # This is a table header
                            table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
                            current_table = doc.add_table(rows=1, cols=len(table_headers))
                            current_table.style = 'Table Grid'
                            
                            # Add headers
                            header_cells = current_table.rows[0].cells
                            for j, header in enumerate(table_headers):
                                header_cells[j].text = header
                                # Make header bold
                                for paragraph in header_cells[j].paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                            
                            in_table = True
                            i += 1  # Skip the separator line
                        else:
                            # Regular paragraph with pipe characters
                            doc.add_paragraph(line)
                    else:
                        # Table data row
                        if line.startswith('|') and len(line.split('|')) > 2:
                            row_cells = [cell.strip() for cell in line.split('|')[1:-1]]
                            if len(row_cells) == len(table_headers):
                                row = current_table.add_row()
                                for j, cell_content in enumerate(row_cells):
                                    row.cells[j].text = cell_content
                            else:
                                # Malformed table row, end table
                                in_table = False
                                current_table = None
                                doc.add_paragraph(line)
                        else:
                            # End of table
                            in_table = False
                            current_table = None
                            doc.add_paragraph(line)
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                else:
                    # Regular paragraph
                    if in_table and not line.startswith('|'):
                        # End table if we encounter non-table content
                        in_table = False
                        current_table = None
                    doc.add_paragraph(line)
                
                i += 1
            
            # Save the document
            doc.save(output_path)
            return str(output_path), "docx"
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            print(f"⚠️  Failed to generate DOCX: {e}")
            print("Saving as markdown instead...")
            # Fallback to markdown
            markdown_path = output_path.with_suffix('.md')
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return str(markdown_path), "md"


def version_callback(value: bool):
    """Callback for version option."""
    if value:
        print(f"TestTeller Agent version: {__version__}")
        raise typer.Exit()


app = typer.Typer(
    help=APP_SHORT_DESCRIPTION + ". Configure the agent via .env file.",
    context_settings={"help_option_names": ["--help", "-h"]})


def get_collection_name(provided_name: str | None = None) -> str:
    """
    Get the collection name to use, with the following priority:
    1. User-provided name
    2. Name from settings
    3. Default fallback name
    """
    if provided_name:
        return provided_name

    default_name = DEFAULT_COLLECTION_NAME

    try:
        if settings and settings.chromadb:
            settings_dict = settings.chromadb.__dict__
            if settings_dict.get('default_collection_name'):
                name = settings_dict['default_collection_name']
                logger.info(
                    "Using default collection name from settings: %s", name)
                return name
    except Exception as e:
        logger.warning("Failed to get collection name from settings: %s", e)

    logger.info("Using fallback default collection name: %s", default_name)
    return default_name


def check_settings():
    """Check if required settings are available and provide guidance if not."""
    if settings is None:
        env_path = os.path.join(os.getcwd(), '.env')
        print("\n⚠️  Configuration Error: Missing or invalid .env file")
        print("\nTo configure TestTeller, you have two options:")
        print("\n1. Run the configuration wizard:")
        print("   testteller configure")
        print("\n2. Manually create a .env file at:")
        print(f"   {env_path}")
        print("\nMinimum required configuration:")
        print('   GOOGLE_API_KEY="your-api-key-here"')
        print("\nFor more information about configuration, visit:")
        print("   https://github.com/yourusername/testteller#configuration")
        raise typer.Exit(code=1)
    return True


def check_api_key_configured():
    """Check if API key is configured and provide user-friendly guidance."""
    if settings is None:
        check_settings()
        return False

    # Reload .env file to ensure latest configuration is loaded
    from .config import load_env
    load_env()

    # Get provider from settings
    provider = settings.llm.provider if settings.llm else 'gemini'
    
    # Check for API key based on provider
    api_key = None
    if settings.api_keys:
        if provider == 'gemini':
            api_key = settings.api_keys.google_api_key
        elif provider == 'openai':
            api_key = settings.api_keys.openai_api_key
        elif provider in ['claude', 'anthropic']:
            api_key = settings.api_keys.claude_api_key

    # Fallback to environment variables
    env_vars = {
        'gemini': 'GOOGLE_API_KEY',
        'openai': 'OPENAI_API_KEY',
        'claude': 'ANTHROPIC_API_KEY',
        'anthropic': 'ANTHROPIC_API_KEY',
        'llama': 'OLLAMA_BASE_URL'  # Llama uses base URL instead of API key
    }
    env_var = env_vars.get(provider, 'GOOGLE_API_KEY')

    if not api_key:
        api_key = os.getenv(env_var)
        
    # Special handling for llama/ollama - check if base URL is configured
    if provider == 'llama' and not api_key:
        # For llama, we don't need an API key, just a base URL
        ollama_url = os.getenv('OLLAMA_BASE_URL')
        if ollama_url:
            return True

    if not api_key:
        print(f"\n⚠️  API Key Required")
        print(f"\nThis command requires an LLM API key to function.")
        print(f"\nYou need to configure your chosen LLM provider's API key first.")
        print("\nTo configure TestTeller, you have two options:")
        print("\n1. Run the configuration wizard (recommended):")
        print("   testteller configure")
        print("\n2. Set the environment variable (e.g.):")
        print(f"   export {env_var}='your-api-key-here'")
        print("\n   Or add it to your .env file (e.g.):")
        print(f'   {env_var}="your-api-key-here"')
        print("\nFor more information about API keys, visit:")
        print("   https://github.com/iAviPro/testteller-agent")
        return False

    return True


def requires_api_key(func):
    """Decorator to ensure API key is configured before running a command."""
    @wraps(func)
    def wrapper(*args, **kwargs):
        if not check_api_key_configured():
            raise typer.Exit(code=1)
        return func(*args, **kwargs)

    @wraps(func)
    async def async_wrapper(*args, **kwargs):
        if not check_api_key_configured():
            raise typer.Exit(code=1)
        return await func(*args, **kwargs)

    # Return the appropriate wrapper based on whether the function is async
    if asyncio.iscoroutinefunction(func):
        return async_wrapper
    else:
        return wrapper


def _get_agent(collection_name: str) -> TestTellerRagAgent:
    check_settings()  # Ensure settings are available

    # Check API key before trying to initialize
    if not check_api_key_configured():
        raise typer.Exit(code=1)

    try:
        return TestTellerRagAgent(collection_name=collection_name)
    except ValueError as e:
        # Handle specific API key errors
        if "API key" in str(e) or "api_key" in str(e).lower():
            check_api_key_configured()  # This will print the helpful message
            raise typer.Exit(code=1)
        else:
            logger.error(
                "Failed to initialize TestCaseAgent for collection '%s': %s", collection_name, e, exc_info=True)
            print(f"\n❌ Error: {e}")
            raise typer.Exit(code=1)
    except Exception as e:
        logger.error(
            "Failed to initialize TestCaseAgent for collection '%s': %s", collection_name, e, exc_info=True)
        print(f"\n❌ Unexpected error: {e}")
        print("\nPlease check the logs for more details.")
        raise typer.Exit(code=1)


async def ingest_docs_async(path: str, collection_name: str, enhanced: bool = True, chunk_size: int = 1000):
    # Disable ChromaDB telemetry to prevent hanging
    os.environ['ANONYMIZED_TELEMETRY'] = 'False'
    os.environ['CHROMA_CLIENT_AUTH_PROVIDER'] = ''
    
    agent = _get_agent(collection_name)

    async def _ingest_task():
        await agent.ingest_documents_from_path(path, enhanced_parsing=enhanced, chunk_size=chunk_size)
        # Force completion of all background operations by getting the count
        # This ensures the vector store has finished processing everything
        count = await agent.get_ingested_data_count()
        # Add a small delay to ensure all vector store operations are truly complete
        await asyncio.sleep(0.5)
        
        # Return both count and success info for display after spinner stops
        return {
            'count': count,
            'enhanced': enhanced,
            'chunk_size': chunk_size,
            'collection_name': collection_name
        }

    # Keep message short to avoid terminal line wrapping
    filename = path.split('/')[-1] if '/' in path else path
    if len(filename) > 40:  # Truncate very long filenames
        filename = filename[:37] + "..."
    
    ingestion_mode = "enhanced" if enhanced else "basic"
    spinner_text = f"Ingesting docs ({ingestion_mode}): {filename}"
    result = await with_spinner(_ingest_task(), spinner_text)

    # Display all completion messages after everything is truly done
    print(
        f"Successfully ingested documents. Collection '{result['collection_name']}' now contains {result['count']} items.")
    if result['enhanced']:
        print(
            f"💡 Enhanced parsing enabled: Documents chunked ({result['chunk_size']} chars) with metadata extraction")
    
    # Force cleanup to prevent hanging
    import gc
    gc.collect()  # Force garbage collection
    await asyncio.sleep(0.1)  # Give time for cleanup


async def ingest_code_async(source_path: str, collection_name: str, no_cleanup_github: bool):
    # Disable ChromaDB telemetry to prevent hanging
    os.environ['ANONYMIZED_TELEMETRY'] = 'False'
    os.environ['CHROMA_CLIENT_AUTH_PROVIDER'] = ''
    
    agent = _get_agent(collection_name)

    async def _ingest_task():
        await agent.ingest_code_from_source(source_path, cleanup_github_after=not no_cleanup_github)
        # Force completion of all background operations by getting the count
        # This ensures the vector store has finished processing everything
        count = await agent.get_ingested_data_count()
        # Add a small delay to ensure all vector store operations are truly complete
        await asyncio.sleep(0.5)
        
        # Return success info for display after spinner stops
        return {
            'count': count,
            'source_path': source_path,
            'collection_name': collection_name
        }

    # Keep message short to avoid terminal line wrapping
    source_name = source_path.split('/')[-1] if '/' in source_path else source_path
    if len(source_name) > 40:  # Truncate very long paths
        source_name = source_name[:37] + "..."
    
    result = await with_spinner(_ingest_task(), f"Ingesting code: {source_name}...")
    print(
        f"Successfully ingested code from '{result['source_path']}'. Collection '{result['collection_name']}' now contains {result['count']} items.")
    
    # Force cleanup to prevent hanging
    import gc
    gc.collect()  # Force garbage collection
    await asyncio.sleep(0.1)  # Give time for cleanup


async def generate_async(query: str, collection_name: str, num_retrieved: int, output_file: str | None, output_format: str = "md"):
    agent = _get_agent(collection_name)
    
    try:
        current_count = await agent.get_ingested_data_count()
        if current_count == 0:
            print(
                f"Warning: Collection '{collection_name}' is empty. Generation will rely on LLM's general knowledge.")
            if not typer.confirm("Proceed anyway?", default=True):
                print("Generation aborted.")
                return

        async def _generate_task():
            test_cases = await agent.generate_test_cases(query, n_retrieved_docs=num_retrieved)
            
            # If feedback is enabled, store the test cases as part of the generation task
            if "Error:" not in test_cases[:20]:
                enable_feedback = os.getenv('ENABLE_TEST_CASE_FEEDBACK', 'true').lower() == 'true'
                if enable_feedback:
                    storage_metadata = {
                        "output_file": output_file if output_file else "none",
                        "num_retrieved_docs": num_retrieved
                    }
                    try:
                        await agent.store_generated_test_cases(test_cases, query, storage_metadata)
                    except Exception as e:
                        logger.error("Failed to store generated test cases: %s", e)
            
            return test_cases

        test_cases = await with_spinner(_generate_task(), f"Generating test cases for query...")
        print("\n--- Generated Test Cases ---")
        print(test_cases)
        print("--- End of Test Cases ---\n")

        if output_file:
            if "Error:" in test_cases[:20]:
                logger.warning(
                    "LLM generation resulted in an error, not saving to file: %s", test_cases)
                print(
                    f"Warning: Test case generation seems to have failed. Not saving to {output_file}.")
            else:
                try:
                    actual_file, actual_format = await save_test_cases_with_format(test_cases, output_file, output_format)
                    print(f"Test cases saved to: {actual_file} (format: {actual_format})")
                except Exception as e:
                    logger.error(
                        "Failed to save test cases to %s: %s", output_file, e, exc_info=True)
                    print(
                        f"Error: Could not save test cases to {output_file}: {e}")
    finally:
        # Always cleanup the agent to prevent hanging
        if hasattr(agent, 'close'):
            agent.close()
        
        # Force cleanup to prevent hanging
        import gc
        gc.collect()
        await asyncio.sleep(0.1)  # Give time for cleanup


async def status_async(collection_name: str):
    """Check status of a collection asynchronously."""
    agent = _get_agent(collection_name)
    count = await agent.get_ingested_data_count()
    print(f"\nCollection '{collection_name}' contains {count} ingested items.")

    # Print ChromaDB connection info
    if agent.vector_store.use_remote:
        print(
            f"ChromaDB connection: Remote at {agent.vector_store.host}:{agent.vector_store.port}")
    else:
        print(f"ChromaDB persistent path: {agent.vector_store.db_path}")


async def clear_data_async(collection_name: str, force: bool):
    if not force:
        confirm = typer.confirm(
            f"Are you sure you want to clear all data from collection '{collection_name}' and remove related cloned repositories?")
        if not confirm:
            print("Operation cancelled.")
            return False  # Return False to indicate cancellation

    # Disable ChromaDB telemetry to prevent hanging
    os.environ['ANONYMIZED_TELEMETRY'] = 'False'
    os.environ['CHROMA_CLIENT_AUTH_PROVIDER'] = ''

    # Import here to avoid circular imports
    from testteller.core.data_ingestion.code_loader import CodeLoader
    import chromadb

    try:
        # Directly clear the collection without initializing the full vector store
        # This avoids the need for LLM/embeddings just to delete data

        # Get ChromaDB configuration from settings
        if settings and settings.chromadb:
            use_remote = settings.chromadb.__dict__.get('use_remote', False)
            persist_directory = settings.chromadb.__dict__.get(
                'persist_directory', 'chroma_db')
            host = settings.chromadb.__dict__.get('host', 'localhost')
            port = settings.chromadb.__dict__.get('port', 8000)
        else:
            use_remote = False
            persist_directory = 'chroma_db'
            host = 'localhost'
            port = 8000

        # Initialize ChromaDB client with telemetry disabled
        if use_remote:
            client = chromadb.HttpClient(host=host, port=port)
        else:
            # Disable telemetry to prevent background threads from hanging the process
            client = chromadb.PersistentClient(
                path=persist_directory,
                settings=chromadb.config.Settings(
                    anonymized_telemetry=False
                )
            )

        # Clear the collection
        async def _clear_task():
            try:
                client.delete_collection(name=collection_name)
                logger.info(f"Deleted collection '{collection_name}'")
            except Exception as e:
                logger.warning(
                    f"Collection '{collection_name}' may not exist: {e}")

            # Also clean up cloned repositories
            code_loader = CodeLoader()
            await code_loader.cleanup_all_repos()
            logger.info("Cleaned up all cloned repositories")

        await with_spinner(_clear_task(), f"Clearing data from collection '{collection_name}'...")
        print(
            f"Successfully cleared data from collection '{collection_name}'.")
        
        # Basic cleanup of ChromaDB client
        try:
            if hasattr(client, 'close'):
                client.close()
        except Exception as cleanup_error:
            logger.debug("Error during client cleanup: %s", cleanup_error)
        
        return True  # Return True to indicate success
    except Exception as e:
        logger.error("Error clearing data: %s", e)
        print(f"Error: Failed to clear data. Details: {e}")
        return False


@app.command()
@requires_api_key
def ingest_docs(
    path: Annotated[str, typer.Argument(help="Path to a document file or directory (supports .md, .txt, .pdf, .docx, .xlsx).")],
    collection_name: Annotated[str, typer.Option(
        "--collection-name", "-c", help="ChromaDB collection name.")] = None,
    enhanced: Annotated[bool, typer.Option(
        "--enhanced", "-e", help="Use enhanced parsing with chunking and metadata extraction")] = True,
    chunk_size: Annotated[int, typer.Option(
        "--chunk-size", "-s", help="Text chunk size for better retrieval (100-5000)")] = 1000
):
    """Ingests documents from a file or directory into a collection."""
    # Get collection name from settings if not provided
    collection_name = get_collection_name(collection_name)

    logger.info("CLI: Ingesting documents from '%s' into collection '%s'",
                path, collection_name)

    if not os.path.exists(path):
        logger.error(
            "Document source path does not exist or is not accessible: %s", path)
        print(
            f"Error: Document source path '{path}' not found or not accessible.")
        raise typer.Exit(code=1)

    # Validate chunk_size
    if not (100 <= chunk_size <= 5000):
        print("❌ Error: chunk-size must be between 100 and 5000 characters")
        raise typer.Exit(code=1)

    # Show ingestion mode
    mode = "enhanced" if enhanced else "basic"
    print(f"\n📄 Document Ingestion ({mode} mode)")
    if enhanced:
        print(f"  • Chunk size: {chunk_size} characters")
        print(f"  • Metadata extraction: enabled")
        print(f"  • Supported formats: .md, .txt, .pdf, .docx, .xlsx")
    else:
        print(f"  • Basic parsing mode")

    try:
        asyncio.run(ingest_docs_async(
            path, collection_name, enhanced, chunk_size))
    except EmbeddingGenerationError as e:
        logger.error(
            "CLI: Embedding generation failed during document ingestion. Error: %s", e, exc_info=True)
        print(f"\n❌ Embedding Generation Failed:")
        print(f"   {e}")
        print("\n💡 Potential Solutions:")
        print("   1. Verify your API key in the .env file is correct and has sufficient quota.")
        print("   2. For Claude, ensure CLAUDE_API_KEY and your selected embedding provider's API key are set.")
        print("   3. For Llama, ensure the Ollama service is running and accessible.")
        print("   4. Check your network connection and firewall settings.")
        print("\nRun 'testteller configure' to re-check your settings.")
        raise typer.Exit(code=1)
    except typer.Exit:
        # Re-raise typer.Exit exceptions to avoid catching them
        raise
    except Exception as e:
        logger.error(
            "CLI: Unhandled error during document ingestion from '%s': %s", path, e, exc_info=True)
        print(f"An unexpected error occurred: {e}")
        raise typer.Exit(code=1)


@app.command()
@requires_api_key
def ingest_code(
    source_path: Annotated[str, typer.Argument(help="URL of the GitHub repository OR path to a local code folder.")],
    collection_name: Annotated[str, typer.Option(
        "--collection-name", "-c", help="ChromaDB collection name.")] = None,
    no_cleanup_github: Annotated[bool, typer.Option(
        "--no-cleanup-github", "-nc", help="Do not delete cloned GitHub repo after ingestion (no effect for local folders).")] = False
):
    """Ingests code from a GitHub repository or local folder into a collection."""
    # Get collection name from settings if not provided
    collection_name = get_collection_name(collection_name)

    logger.info("CLI: Ingesting code from '%s' into collection '%s'",
                source_path, collection_name)

    # For local paths, check if they exist
    if not source_path.startswith(('http://', 'https://', 'git@')) and not os.path.exists(source_path):
        logger.error(
            "Local source path does not exist or is not accessible: %s", source_path)
        print(
            f"Error: Local source path '{source_path}' not found or not accessible.")
        raise typer.Exit(code=1)

    try:
        asyncio.run(ingest_code_async(
            source_path, collection_name, no_cleanup_github))
    except EmbeddingGenerationError as e:
        logger.error(
            "CLI: Embedding generation failed during code ingestion. Error: %s", e, exc_info=True)
        print(f"\n❌ Embedding Generation Failed:")
        print(f"   {e}")
        print("\n💡 Potential Solutions:")
        print("   1. Verify your API key in the .env file is correct and has sufficient quota.")
        print("   2. For Claude, ensure CLAUDE_API_KEY and your selected embedding provider's API key are set.")
        print("   3. For Llama, ensure the Ollama service is running and accessible.")
        print("   4. Check your network connection and firewall settings.")
        print("\nRun 'testteller configure' to re-check your settings.")
        raise typer.Exit(code=1)
    except typer.Exit:
        # Re-raise typer.Exit exceptions to avoid catching them
        raise
    except Exception as e:
        logger.error(
            "CLI: Unhandled error during code ingestion from '%s': %s", source_path, e, exc_info=True)
        print(f"An unexpected error occurred: {e}")
        raise typer.Exit(code=1)


@app.command()
@requires_api_key
def generate(
    query: Annotated[str, typer.Argument(help="Query for test case generation.")],
    collection_name: Annotated[str, typer.Option(
        "--collection-name", "-c", help="ChromaDB collection name.")] = None,
    num_retrieved: Annotated[int, typer.Option(
        "--num-retrieved", "-n", min=0, max=20, help="Number of docs for context.")] = 5,
    output_file: Annotated[str, typer.Option(
        "--output-file", "-o", help=f"Optional: Save test cases to this file. If not provided, uses OUTPUT_FILE_PATH from .env or defaults to {DEFAULT_OUTPUT_FILE}")] = None,
    output_format: Annotated[str, typer.Option(
        "--output-format", "-f", help="Output format for test cases: md, pdf, docx. Uses TEST_OUTPUT_FORMAT from .env if not specified.")] = None
):
    """Generates test cases based on query and knowledge base."""
    logger.info(
        "CLI: Generating test cases for query: '%s...', Collection: %s", query[:50], collection_name)

    # Get collection name from settings if not provided
    collection_name = get_collection_name(collection_name)
    
    # Validate and get output format
    final_output_format = output_format
    if not final_output_format:
        # Try to get from environment/settings
        final_output_format = os.getenv('TEST_OUTPUT_FORMAT', DEFAULT_TEST_OUTPUT_FORMAT)
        logger.info("Using output format from environment/default: %s", final_output_format)
    
    if final_output_format not in SUPPORTED_TEST_OUTPUT_FORMATS:
        print(f"❌ Error: Unsupported output format '{final_output_format}'")
        print(f"Supported formats: {', '.join(SUPPORTED_TEST_OUTPUT_FORMATS)}")
        raise typer.Exit(code=1)

    # Determine output file path with new directory structure and format
    final_output_file = output_file
    if not final_output_file:
        # Generate timestamped filename in the new directory structure
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"testteller-testcases-{timestamp}"
        
        try:
            if settings and settings.output:
                settings_dict = settings.output.__dict__
                if settings_dict.get('output_file_path'):
                    # Use configured path but update extension based on format
                    configured_path = Path(settings_dict['output_file_path'])
                    final_output_file = str(configured_path.parent / f"{configured_path.stem}-{timestamp}.{final_output_format}")
                    logger.info("Using configured output path with timestamp: %s", final_output_file)
        except Exception as e:
            logger.warning("Failed to get output file path from settings: %s", e)

        if not final_output_file:
            # Use new default directory structure
            output_dir = Path(DEFAULT_TEST_GENERATION_DIR)
            output_dir.mkdir(exist_ok=True)
            final_output_file = str(output_dir / f"{base_filename}.{final_output_format}")
            logger.info("Using default output path with timestamp: %s", final_output_file)
    else:
        # If user provided output file, ensure the directory exists
        output_path = Path(final_output_file)
        output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        asyncio.run(generate_async(
            query, collection_name, num_retrieved, final_output_file, final_output_format))
    except typer.Exit:
        # Re-raise typer.Exit exceptions to avoid catching them
        raise
    except Exception as e:
        logger.error(
            "CLI: Unhandled error during test case generation: %s", e, exc_info=True)
        print(f"An unexpected error occurred: {e}")
        raise typer.Exit(code=1)


@app.command()
@requires_api_key
def status(
    collection_name: Annotated[str, typer.Option(
        "--collection-name", "-c", help="ChromaDB collection name.")] = None
):
    """Checks status of a collection."""
    # Get collection name from settings if not provided
    collection_name = get_collection_name(collection_name)

    logger.info("CLI: Checking status for collection: %s", collection_name)
    try:
        asyncio.run(status_async(collection_name))
    except typer.Exit:
        # Re-raise typer.Exit exceptions to avoid catching them
        raise
    except Exception as e:
        logger.error(
            "CLI: Unhandled error during status check: %s", e, exc_info=True)
        print(f"An unexpected error occurred: {e}")
        raise typer.Exit(code=1)


@app.command()
def clear_data(
    collection_name: Annotated[str, typer.Option(
        "--collection-name", "-c", help="ChromaDB collection to clear.")] = None,
    force: Annotated[bool, typer.Option(
        "--force", "-f", help="Force clear without confirmation.")] = False
):
    """Clears ingested data from a collection."""
    # Get collection name from settings if not provided
    collection_name = get_collection_name(collection_name)

    logger.info("CLI: Clearing data for collection: %s", collection_name)
    try:
        result = asyncio.run(clear_data_async(collection_name, force))
        if result is False:
            # Operation was cancelled by user
            os._exit(0)
        # Force exit after successful completion to avoid hanging on background threads
        os._exit(0)
    except typer.Exit:
        # Re-raise typer.Exit exceptions to avoid catching them
        raise
    except Exception as e:
        logger.error(
            "CLI: Unhandled error during data clearing: %s", e, exc_info=True)
        print(f"An unexpected error occurred: {e}")
        raise typer.Exit(code=1)


@app.command()
def configure(
    provider: Annotated[str, typer.Option(
        "--provider", "-p", help="Quick setup for specific provider (gemini, openai, claude, llama)")] = None,
    automator_agent: Annotated[bool, typer.Option(
        "--automator-agent", "-aa", help="Configure Automator Agent settings only")] = False
):
    """Interactive configuration wizard to set up TestTeller."""
    from testteller.core.config import ConfigurationWizard, run_provider_only_setup, run_automation_only_setup
    from testteller.core.config import UIMode

    env_path = Path.cwd() / ".env"

    try:
        # Handle different configuration modes
        if automator_agent:
            # Configure Automator Agent settings only
            success = run_automation_only_setup(UIMode.CLI)
            if not success:
                # Automator Agent configuration failed (not user cancellation)
                raise typer.Exit(code=1)
            return

        elif provider:
            # Quick setup for specific provider
            if provider not in SUPPORTED_LLM_PROVIDERS:
                print(f"❌ Unsupported provider: {provider}")
                print(
                    f"Supported providers: {', '.join(SUPPORTED_LLM_PROVIDERS)}")
                raise typer.Exit(code=1)

            success = run_provider_only_setup(provider, UIMode.CLI)
            if not success:
                # Provider configuration failed
                raise typer.Exit(code=1)
            return

        # Full configuration wizard
        wizard = ConfigurationWizard(UIMode.CLI)
        success = wizard.run(env_path)

        if success:
            print("\n🚀 TestTeller is now ready to use!")
            print("\n📚 Next steps:")
            print("  testteller --help                    # View all commands")
            print("  testteller ingest-docs <path>        # Add documents")
            print("  testteller ingest-code <repo_url>    # Add code")
            print("  testteller generate \"<query>\"        # Generate tests")
            if HAS_AUTOMATION:
                print(
                    "  testteller automate test-cases.md    # Generate automation code")
        else:
            # Configuration was cancelled by user, not a failure
            raise typer.Exit(code=0)

    except KeyboardInterrupt:
        print("\n⚠️  Configuration cancelled by user.")
        raise typer.Exit(code=1)
    except typer.Exit:
        # Re-raise typer.Exit without logging as error
        raise
    except Exception as e:
        logger.error("Configuration wizard failed: %s", e, exc_info=True)
        print("❌ Configuration failed.")
        raise typer.Exit(code=1)


# TestTeller automation command (if available)
if HAS_AUTOMATION:
    @app.command()
    @requires_api_key
    def automate(
        input_file: Annotated[str, typer.Argument(help="Path to test cases file (supports .md, .txt, .pdf, .docx, .xlsx)")],
        collection_name: Annotated[str, typer.Option(
            "--collection-name", "-c", help="ChromaDB collection name for application context")] = None,
        language: Annotated[str, typer.Option(
            "--language", "-l", help="Programming language for test automation")] = None,
        framework: Annotated[str, typer.Option(
            "--framework", "-F", help="Test framework to use")] = None,
        output_dir: Annotated[str, typer.Option(
            "--output-dir", "-o", help="Output directory for generated tests")] = "./testteller_automated_tests",
        interactive: Annotated[bool, typer.Option(
            "--interactive", "-i", help="Interactive mode to select test cases")] = False,
        num_context_docs: Annotated[int, typer.Option(
            "--num-context", "-n", min=1, max=20, help="Number of context documents to retrieve")] = 5,
        verbose: Annotated[bool, typer.Option(
            "--verbose", "-v", help="Enable verbose logging")] = False
    ):
        """Generate automation code using RAG-enhanced approach with vector store knowledge."""
        if not HAS_AUTOMATION:
            print(
                "❌ Automation functionality not available. Please check your installation.")
            raise typer.Exit(code=1)

        # Delegate to the consolidated automate_command with correct parameters
        automate_command(
            input_file=input_file,
            collection_name=collection_name,
            language=language,
            framework=framework,
            output_dir=output_dir,
            interactive=interactive,
            num_context_docs=num_context_docs,
            verbose=verbose
        )


@app.callback()
def main(
    _: Annotated[bool, typer.Option(
        "--version", "-v",
        help="Show version and exit",
        callback=version_callback,
        is_eager=True
    )] = False
):
    """Complete AI Test Agent for Generation and Automation. Configure the agent via your .env file."""
    pass


def app_runner():
    """
    This function is the entry point for the CLI script defined in pyproject.toml.
    It ensures logging is set up and then runs the Typer application.
    """
    try:
        app()
    except Exception as e:
        logger.error("Unhandled error in CLI: %s", e, exc_info=True)
        print(f"\n❌ An unexpected error occurred: {e}")
        raise typer.Exit(code=1)


if __name__ == "__main__":
    app_runner()
