from enum import Enum

from docx import Document
from docx.styles.styles import Styles
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from py_docx_creator.AbstractClasses import DocumentCreator, DocumentWriter, DocumentStyle, PageStyle, ParagraphStyle, TextStyle


class FontNames(Enum):
    """Перечень наименований шрифтов"""
    TimesNewRoman = "Times New Roman"
class DocumentStyles(Enum):
    """Перечень стилей документа"""
    Normal = "Normal"
class CoreDocumentCreator(DocumentCreator):
    
    def __init__(self):
        super().__init__()
        self._document = None
        self._file_name = None
    def create_document(self, file_name):
        self.document = Document()
        self.file_name = file_name

    def load_document(self):
        self.document = Document(self.path_to_document or self.file_name)
        
    def save_document(self):
        self.document.save(self.file_name)
class CoreDocumentWriter(DocumentWriter):
    
    def __init__(self):
        super().__init__()
    
    @staticmethod
    def add_paragraph_to_document(document: Document) -> Paragraph:
        return document.add_paragraph()
    
    @staticmethod
    def add_run_to_paragraph(paragraph: Paragraph, text: str) -> Run:
        return paragraph.add_run(text)
    
    @staticmethod
    def add_page_break(document: Document) -> None:
        document.add_page_break()
class CoreDocumentStyle(DocumentStyle):

    def __init__(self):
        super().__init__()

    def get_document_style(self, document: Document) -> Styles:
        return document.style[f"{self.document_style}"]
class CorePageStyle(PageStyle):
    """Абстрактный класс реализации применения стиля к страницам документа"""

    def __init__(self):
        super().__init__()

    def apply_style(self, document):
        for section in document.sections:
            section.top_margin = self.top_margin
            section.bottom_margin = self.bottom_margin
            section.left_margin = self.left_margin
            section.right_margin = self.right_margin
class CoreParagraphFormat(ParagraphStyle):
    """ Абстрактный класс реализации метода применения стиля к параграфу"""

    def __init__(self):
        super().__init__()

    def apply_style(self, paragraph) -> None:
        paragraph_style = paragraph.paragraph_format

        if self.paragraph_alignment is not None:
            paragraph_style.alignment = self.paragraph_alignment

        if self.space_after_paragraph is not None:
            paragraph_style.space_after = self.space_after_paragraph

        if self.paragraph_left_indent is not None:
            paragraph_style.left_indent = self.paragraph_left_indent

        if self.paragraph_right_indent is not None:
            paragraph_style.right_indent = self.paragraph_right_indent

        if self.paragraph_line_spacing is not None:
            paragraph_style.line_spacing = self.paragraph_line_spacing

        if self.paragraph_first_line_indent is not None:
            paragraph_style.first_line_indent = self.paragraph_first_line_indent
class CoreTextStyle(TextStyle):
    """Основной стиль текста"""    
    def __init__(self):
        super().__init__()

    
    def apply_style(self, run):
        if self.text_bold is not None:
            run.font.bold = self.text_bold
        if self.text_italic is not None:
            run.font.italic = self.text_italic
        if self.text_underline is not None:
            run.font.underline = self.text_underline
        if self.font_size is not None:
            run.font.size = self.font_size
        if self.font_name is not None:
            run.font.name = self.font_name






