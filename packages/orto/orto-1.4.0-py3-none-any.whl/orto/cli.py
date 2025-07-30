import argparse
import xyz_py as xyzp
from xyz_py.atomic import elements as atomic_elements
import sys
import pathlib
import os
import copy
import subprocess
import csv
import numpy as np
import re
from mmap import mmap, ACCESS_READ
from shutil import move as shutilmove

from . import job
from . import utils as ut
from . import constants as cst
from .exceptions import DataNotFoundError, DataFormattingError

_SHOW_CONV = {
    'on': True,
    'save': False,
    'show': True,
    'off': False
}

_SAVE_CONV = {
    'on': True,
    'save': True,
    'show': False,
    'off': False
}


def extract_coords_func(uargs, save=True):
    '''
    Wrapper for extract_coords function

    Parameters
    ----------
    args : argparser object
        command line arguments
    save : bool, default=True
        If True, saves data to file. If False, prints to stdout.

    Returns
    -------
        None
    '''
    from . import extractor as oe

    # Open file and extract coordinates
    labels, coords = oe.get_coords(
        uargs.output_file,
        coord_type=uargs.type,
        index_style=uargs.index_style
    )

    if save:
        # Save to new .xyz file
        xyzp.save_xyz(
            f'{uargs.output_file.stem}_coords.xyz',
            labels,
            coords,
            comment=f'Coordinates extracted from {uargs.output_file}'
        )
    else:
        for lbl, crd in zip(labels, coords):
            print(f'{lbl} {crd[0]:.4f} {crd[1]:.4f} {crd[2]:.4f}')

    return


def extract_sf_energies_func(uargs):
    '''
    Wrapper for cli call to extract spin-free energies
    '''
    from . import extractor as oe
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt

    all_data = oe.SpinFreeEnergyExtractor().extract(
        uargs.output_file
    )

    if uargs.output_format == 'txt':
        out_name = f'{uargs.output_file.stem}_sfenergies.txt'
        with open(out_name, 'w') as f:
            for it, data in enumerate(all_data):
                if len(all_data) > 1:
                    f.write(f'# Section {it + 1:d}\n')

                f.write('State, Root, Multiplicity, Relative Energy (cm⁻¹)\n')

                f.write(
                    f'1, {data['root'][0]:d}, {data['multiplicity'][0]:d}, 0\n' # noqa
                )
                for rit in range(1, len(data['root'])):
                    f.write(
                        f'{rit + 1:d}, {data['root'][rit]:d}, {data['multiplicity'][rit]:d}, {data['delta energy (cm^-1)'][rit - 1]:.2f}\n' # noqa
                    )
                if len(all_data) > 1:
                    f.write('-------------------\n')

    if uargs.output_format == 'docx':
        out_name = f'{uargs.output_file.stem}_sfenergies.docx'

        title = 'Spin-Free energies'

        # Create document
        doc = Document()

        doc.add_heading(title, 0)

        # Add style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9)

        # For each extracted section, print matrix, vectors, and values
        for it, data in enumerate(all_data):
            if len(all_data) > 1:
                doc.add_paragraph(f'\nSection {it + 1:d}\n')

            # Table of data
            table = doc.add_table(rows=len(data['root']) + 1, cols=4)

            table.cell(0, 0).text = 'State'
            table.cell(0, 1).text = 'Root'
            table.cell(0, 2).text = 'Multiplicity'
            table.cell(0, 3).text = 'Relative Energy (cm'
            ss = table.cell(0, 3).paragraphs[0].add_run('-1')
            ss.font.superscript = True
            table.cell(0, 3).paragraphs[0].add_run(')')

            # Add data
            for rit in range(len(data['root'])):
                table.cell(rit + 1, 0).text = f'{rit + 1:d}'
                table.cell(rit + 1, 1).text = f'{data['root'][rit]:d}'
                table.cell(rit + 1, 2).text = f'{data['multiplicity'][rit]:d}'
                if rit == 0:
                    table.cell(1, 3).text = '0'
                else:
                    table.cell(rit + 1, 3).text = f'{data['delta energy (cm^-1)'][rit - 1]:.2f}' # noqa

            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # noqa
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].style = 'Normal'

        doc.save(out_name)

    ut.cprint(f'Data written to {out_name}', 'cyan')

    return


def extract_so_energies_func(uargs):
    '''
    Wrapper for cli call to extract spin-orbit energies
    '''
    from . import extractor as oe

    all_energies = oe.SpinOrbitEnergyExtractor().extract(
        uargs.output_file
    )

    out_name = f'{uargs.output_file.stem}_soenergies.txt'
    with open(out_name, 'w') as f:
        for it, energies in enumerate(all_energies):
            if len(all_energies) > 1:
                f.write(f'# Section {it + 1:d}\n')

            f.write('State, Relative Energy (cm⁻¹)\n')

            for eit, energy in enumerate(energies):
                f.write(
                    f'{eit + 1:d}, {energy:.4f}\n'
                )
            if len(all_energies) > 1:
                f.write('-------------------\n')

    ut.cprint(f'Data written to {out_name}', 'cyan')

    return


def extract_gmatrix_func(uargs, save=True):
    '''
    Wrapper for cli call to extract gmatrix

    Parameters
    ----------
    args : argparser object
        command line arguments
    save : bool, default=True
        If True, saves data to file. If False, prints to stdout.

    Returns
    -------
        None
    '''
    from . import extractor as oe
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt

    choices = {
        'total': oe.GMatrixExtractor,
        'L': oe.GMatrixLExtractor,
        'S': oe.GMatrixSExtractor,
        'eff': oe.GMatrixEffectiveExtractor,
    }

    all_data = choices[uargs.type]().extract(uargs.output_file)

    if not save:
        for it, data in enumerate(all_data):
            print(f'Section {it + 1:d}')
            for key, val in data.items():
                print(f'{key}:')
                print(val)
        sys.exit(0)

    if uargs.output_format == 'txt':
        out_name = f'{uargs.output_file.stem}_gmatrix.txt'
        with open(out_name, 'w') as f:
            for it, data in enumerate(all_data):
                if len(all_data) > 1:
                    f.write(f'Section {it + 1:d}\n')
                for key, val in data.items():
                    f.write(f'{key}:\n')
                    f.write(str(val).replace('[', '').replace(']', ''))
                    f.write('\n')

    if uargs.output_format == 'docx':
        out_name = f'{uargs.output_file.stem}_gmatrix.docx'

        titles = {
            'total': 'ELECTRONIC G-MATRIX',
            'L': 'ELECTRONIC G-MATRIX: L contribution',
            'S': 'ELECTRONIC G-MATRIX: S contribution',
            'eff': 'ELECTRONIC G-MATRIX FROM EFFECTIVE HAMILTONIAN'
        }

        title = titles[uargs.type]

        # Create document
        doc = Document()

        doc.add_heading(title, 0)

        # Add style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9)

        # For each extracted section, print matrix, vectors, and values
        for it, data in enumerate(all_data):
            if len(all_data) > 1:
                doc.add_paragraph(f'Section {it + 1:d}')

            if uargs.type == 'eff':
                doc.add_paragraph(f'Effective S={data['spin_mult']:2d}')

            # Full matrix
            matrix = doc.add_table(rows=3, cols=3)

            matrix.cell(0, 0).text = '{:.4f}'.format(data['matrix'][0, 0])
            matrix.cell(1, 0).text = '{:.4f}'.format(data['matrix'][0, 1])
            matrix.cell(2, 0).text = '{:.4f}'.format(data['matrix'][0, 2])

            matrix.cell(0, 1).text = '{:.4f}'.format(data['matrix'][1, 0])
            matrix.cell(1, 1).text = '{:.4f}'.format(data['matrix'][1, 1])
            matrix.cell(2, 1).text = '{:.4f}'.format(data['matrix'][1, 2])

            matrix.cell(0, 2).text = '{:.4f}'.format(data['matrix'][2, 0])
            matrix.cell(1, 2).text = '{:.4f}'.format(data['matrix'][2, 1])
            matrix.cell(2, 2).text = '{:.4f}'.format(data['matrix'][2, 2])

            doc.add_paragraph('\n')

            for row in matrix.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # noqa
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].style = 'Normal'

            # g values and g vectors
            table = doc.add_table(rows=4, cols=4)
            table.cell(0, 1).merge(table.cell(0, 1)).merge(table.cell(0, 2)).merge(table.cell(0, 3)) # noqa

            table.cell(0, 0).text = 'Values'
            table.cell(0, 1).text = 'Vectors'

            table.cell(1, 0).text = '{:.4f}'.format(data['values'][0])
            table.cell(2, 0).text = '{:.4f}'.format(data['values'][1])
            table.cell(3, 0).text = '{:.4f}'.format(data['values'][2])

            table.cell(1, 1).text = '{:.4f}'.format(data['vectors'][0, 0])
            table.cell(2, 1).text = '{:.4f}'.format(data['vectors'][0, 1])
            table.cell(3, 1).text = '{:.4f}'.format(data['vectors'][0, 2])

            table.cell(1, 2).text = '{:.4f}'.format(data['vectors'][1, 0])
            table.cell(2, 2).text = '{:.4f}'.format(data['vectors'][1, 1])
            table.cell(3, 2).text = '{:.4f}'.format(data['vectors'][1, 2])

            table.cell(1, 3).text = '{:.4f}'.format(data['vectors'][2, 0])
            table.cell(2, 3).text = '{:.4f}'.format(data['vectors'][2, 1])
            table.cell(3, 3).text = '{:.4f}'.format(data['vectors'][2, 2])

            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # noqa
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].style = 'Normal'

        doc.save(out_name)

    ut.cprint(f'Data written to {out_name}', 'cyan')

    return


def gen_trunc_molden_func(uargs):
    '''
    Wrapper for CLI gen truncmolden call

    Parameters
    ----------
    uargs : argparser object
        User arguments

    Returns
    -------
    None
    '''

    # Read molden file in as binary string
    # and find number of MOs by counting number of
    # occurrences of 'Occup='
    _patt = re.compile(b'Sym=')
    with open(uargs.input_file, mode="r") as file_obj:
        with mmap(file_obj.fileno(), length=0, access=ACCESS_READ) as mmap_obj:
            n_MO = len(_patt.findall(mmap_obj))

    _patt = re.compile(b'Occup= 0.000000')
    with open(uargs.input_file, mode="r") as file_obj:
        with mmap(file_obj.fileno(), length=0, access=ACCESS_READ) as mmap_obj:
            n_virt = len(_patt.findall(mmap_obj))

    ut.cprint(
        (
            f'{n_virt}/{n_MO} MOs are virtual...'
        ),
        'cyan'
    )

    if uargs.output_file is None:
        # If no output file specified, use input file name
        # with .molden extension
        uargs.output_file = '.tmp.molden'

    # Trim file
    _count = 0
    final = False
    with open(uargs.input_file, mode="r") as old:
        with open(uargs.output_file, mode="w") as new:
            # Read in molden file line by line
            for line in old:
                if 'Occup= 0.000000' in line:
                    _count += 1
                if _count == uargs.n_virt:
                    final = True
                if line.startswith(f'{n_MO:d}') and final:
                    new.write(line)
                    break
                else:
                    new.write(line)

    ut.cprint(f'... trimming to {uargs.n_virt} virtual orbitals\n', 'cyan')

    # If no output file given
    if uargs.output_file == '.tmp.molden':
        # Copy new file to original name
        shutilmove(uargs.output_file, uargs.input_file)
        uargs.output_file = uargs.input_file

    ut.cprint(f'New molden file written to {uargs.output_file}', 'cyan')

    return


def gen_job_func(uargs):
    '''
    Wrapper for CLI gen job call

    Parameters
    ----------
    uargs : argparser object
        User arguments

    Returns
    -------
    None
    '''
    from . import input as inp

    for input_file in uargs.input_files:

        # Check input exists
        if not input_file.exists:
            ut.red_exit('Cannot locate {}'.format(input_file.name))

        oj = job.OrcaJob(
            input_file
        )

        # Get orca module load command
        orca_args = [
            'orca_load'
        ]

        required = [
            'orca_load'
        ]

        for oarg in orca_args:
            uarg_val = getattr(uargs, oarg)
            if len(uarg_val):
                oarg_val = copy.copy(uarg_val)
            elif os.getenv(f'orto_{oarg}'):
                try:
                    if len(os.getenv(f'orto_{oarg}')):
                        oarg_val = os.getenv(f'orto_{oarg}')
                except ValueError:
                    ut.red_exit(
                        (
                            f'Error in orto_{oarg} environment variable'
                        )
                    )
            elif oarg in required:
                ut.red_exit(
                    (
                        f'Missing orto_{oarg} environment variable or '
                        f'--{oarg} argument'
                    )
                )
            else:
                oarg_val = ''

            if oarg == 'orca_load':
                oarg = 'load'
                if 'module load' not in oarg_val:
                    oarg_val = f'module load {oarg_val}'

            setattr(oj, oarg, oarg_val)

        # Check xyz file is present
        try:
            inp.check_coord(oj.input_file, uargs.skip_xyz)
        except (DataNotFoundError, DataFormattingError) as e:
            ut.red_exit(str(e))

        # Check for moread and moinp
        try:
            inp.check_moinp_moread(oj.input_file)
        except (DataNotFoundError, DataFormattingError) as e:
            ut.red_exit(str(e))

        # Submitter configuration options
        # currently hardcoded for slurm!
        config = {}

        # Get nprocs and maxcore
        try:
            n_procs = inp.get_nprocs(oj.input_file)
            maxcore = inp.get_maxcore(oj.input_file)
        except (DataNotFoundError, DataFormattingError) as e:
            ut.red_exit(str(e))

        # If memory and procs specified as arguments, give warning when
        # these are smaller than the number in the input file
        if uargs.n_procs:
            if n_procs > uargs.n_procs:
                ut.red_exit('Too few processors requested for input file')
            # Use cli value
            config['ntasks_per_node'] = uargs.n_procs
        else:
            # Use orca file value
            config['ntasks_per_node'] = n_procs

        if uargs.memory:
            if uargs.memory * uargs.n_procs < n_procs * maxcore:
                ut.red_exit('Requested too little memory for orca input')
            config['mem_per_cpu'] = uargs.memory
        else:
            # Use orca file value
            config['mem_per_cpu'] = maxcore

        # Check if NBO is requested
        if inp.get_nbo(oj.input_file):
            # Check if NBO module has been provided to orto
            try:
                if os.getenv('orto_nbo_load') is not None:
                    nbo_module = os.getenv('orto_nbo_load')
                else:
                    ut.red_exit(
                        'Missing orto_nbo_load environment variable'
                    )
            except ValueError:
                ut.red_exit(
                    (
                        'Missing or malformed orto_nbo_load'
                        'environment variable'
                    )
                )
            oj.pre_orca += f'module load {nbo_module}\n'
            oj.pre_orca += f'export NBOFIL={oj.input_file.stem}\n'
        else:
            nbo_module = None

        # Set SLURM error and output file names
        config['error'] = 'slurm.%j.e'
        config['output'] = 'slurm.%j.o'

        # Add call to orca_2mkl to create molden file from gbw
        if not uargs.no_molden:
            oj.post_orca += 'orca_2mkl {} -molden'.format(oj.input_file.stem)
            oj.post_orca += '\norto gen trunc_molden {}.molden.input'.format(
                oj.input_file.stem
            )

        # Write job script
        # with submitter configuration options specified
        oj.write_script(True, **config)

        # Submit to queue
        if not uargs.no_sub:
            subprocess.call(
                'cd {}; {} "{}"; cd ../'.format(
                    oj.input_file.parents[0],
                    oj.Job.SUBMIT_COMMAND,
                    oj.job_file
                    ),
                shell=True
            )
    return


def plot_abs_func(uargs):
    '''
    Wrapper for CLI plot abs call

    Parameters
    ----------
    uargs : argparser object
        User arguments

    Returns
    -------
    None
    '''
    import matplotlib.pyplot as plt
    import matplotlib as mpl
    from . import plotter
    from . import extractor as oe

    # Set user specified font name
    if os.getenv('orto_fontname'):
        try:
            plt.rcParams['font.family'] = os.getenv('orto_fontname')
        except ValueError:
            ut.cprint('Error in orto_fontname environment variable', 'red')
            sys.exit(1)

    # Change matplotlib font size to be larger
    mpl.rcParams.update({'font.size': 12})

    version = oe.OrcaVersionExtractor.extract(uargs.output_file)

    if not len(version):
        ut.cprint(
            'Warning: Cannot find version number in Orca output file',
            'black_yellowbg'
        )
        version = [6, 0, 0]

    if version[0] < 6:
        if uargs.intensities == 'electric':
            all_data = oe.OldAbsorptionElectricDipoleExtractor.extract(
                uargs.output_file
            )
        elif uargs.intensities == 'velocity':
            all_data = oe.OldAbsorptionVelocityDipoleExtractor.extract(
                uargs.output_file
            )
    elif version[0] >= 6:
        if uargs.intensities == 'electric':
            all_data = oe.AbsorptionElectricDipoleExtractor.extract(
                uargs.output_file
            )
        elif uargs.intensities == 'velocity':
            all_data = oe.AbsorptionVelocityDipoleExtractor.extract(
                uargs.output_file
            )
        elif uargs.intensities == 'semi-classical':
            all_data = oe.AbsorptionSemiClassicalDipoleExtractor.extract(
                uargs.output_file
            )

    ut.cprint('Using intensities: {}'.format(uargs.intensities), 'cyan')

    # Plot each section
    for it, data in enumerate(all_data):

        if len(all_data) > 1:
            save_name = f'absorption_spectrum_section_{it:d}.png'
        else:
            save_name = 'absorption_spectrum.png'

        if uargs.x_unit == 'wavenumber':
            x_values = data['energy (cm^-1)']
        elif uargs.x_unit == 'wavelength':
            x_values = 1E7 / data['energy (cm^-1)']
        elif uargs.x_unit == 'energy':
            x_values = data['energy (ev)']

        # Plot absorption spectrum
        fig, ax = plotter.plot_abs(
            x_values,
            uargs.x_unit,
            data['fosc'],
            show=_SHOW_CONV[uargs.plot],
            save=_SAVE_CONV[uargs.plot],
            save_name=save_name,
            x_lim=uargs.x_lim,
            y_lim=uargs.y_lim,
            x_shift=uargs.shift,
            linewidth=uargs.linewidth,
            lineshape=uargs.lineshape,
            window_title=f'Absorption Spectrum from {uargs.output_file}',
            osc_style=uargs.osc_style,
            normalise=uargs.normalise_absorption
        )

        if uargs.x_unit == 'wavenumber':
            ax[0].set_xlim([0, 50000])
        plt.show()

    return


def plot_ir_func(uargs):
    '''
    Wrapper for CLI plot_ir call

    Parameters
    ----------
    uargs : argparser object
        User arguments

    Returns
    -------
    None
    '''
    import matplotlib.pyplot as plt
    import matplotlib as mpl
    from . import plotter
    from . import extractor as oe

    # Set user specified font name
    if os.getenv('orto_fontname'):
        try:
            plt.rcParams['font.family'] = os.getenv('orto_fontname')
        except ValueError:
            ut.cprint('Error in orto_fontname environment variable', 'red')
            sys.exit(1)

    # Change matplotlib font size to be larger
    mpl.rcParams.update({'font.size': 12})

    # Extract frequency information
    data = oe.FrequencyExtractor.extract(uargs.output_file)

    if not len(data):
        ut.red_exit(f'Cannot find frequencies in file {uargs.output_file}')

    data = data[0]

    # Plot infrared spectrum
    plotter.plot_ir(
        data['energy (cm^-1)'],
        data['epsilon (L mol^-1 cm^-1)'],
        linewidth=uargs.linewidth,
        lineshape=uargs.lineshape,
        window_title=f'Infrared Spectrum from {uargs.output_file}',
        show=True
    )

    return


def distort_func(uargs):
    '''
    Distorts molecule along specified normal mode

    Parameters
    ----------
    args : argparser object
        command line arguments

    Returns
    -------
        None

    '''
    from . import extractor as oe

    # Open file and extract coordinates
    labels, coords = oe.get_coords(
        uargs.output_file
    )

    # Extract frequency information
    data = oe.FrequencyExtractor.extract(uargs.output_file)

    ut.cprint(
        'Distorting along mode #{}:  {: .2f} cm⁻¹'.format(
            uargs.mode_number,
            data[0]['energy (cm^-1)'][uargs.mode_number]
            ),
        'cyan'
        )

    coords += uargs.scale * data[0]['displacements'][:, uargs.mode_number]

    comment = (
        f'Coordinates from {uargs.output_file} distorted by {uargs.scale:f} unit of' # noqa
        f' Mode #{uargs.mode_number}'
    )

    labels_nn = xyzp.remove_label_indices(labels)

    xyzp.save_xyz('distorted.xyz', labels_nn, coords, comment=comment)

    return


def extract_orbs_func(uargs, save=True) -> None:
    '''
    Extracts Loewdin Orbital contributions from orca output file

    Parameters
    ----------
    args : argparser object
        command line arguments
    save : bool, default=True
        If True, saves data to file. If False, prints to stdout.

    Returns
    -------
        None
    '''
    from . import extractor as oe

    # Check for spin in file, if present
    try:
        mult = oe.MultiplicityInputExtractor.extract(uargs.output_file)[0]
        # Disable spin if not present in file
        if mult == 1:
            uargs.spin = None
    except DataNotFoundError:
        pass

    extractors = {
        'orb_comp': lambda x: oe.LoewdinCompositionExtractor.extract(x),
        'redorb_pop': lambda x: oe.LoewdinReducedOrbitalPopulationExtractor.extract( # noqa
            x, spin=uargs.spin
        ),
        'orb_pop': lambda x: oe.LoewdinOrbitalPopulationExtractor.extract(
            x, spin=uargs.spin
        )
    }

    extractor_names = {
        'orb_comp': 'Loewdin Orbital Composition',
        'orb_pop': 'Loewdin Orbital Population',
        'redorb_pop': 'Loewdin Reduced Orbital Population'
    }

    if uargs.flavour == 'first_match':
        failed = 0
        for name, extractor in extractors.items():
            try:
                data = extractor(
                    uargs.output_file
                )
                uargs.flavour = name
                break
            except DataNotFoundError:
                failed += 1
        if failed == len(extractors):
            ut.red_exit(
                'Cannot find Loewdin orbital contributions in file'
            )
    else:
        try:
            data = extractors[uargs.flavour](
                uargs.output_file
            )
        except DataNotFoundError as dne:
            ut.red_exit(str(dne))

    # Unpack data
    contributions, occupancies, energies = data[0]

    # Trim contributions to selected range of MOs
    if uargs.active:
        keep = [
            it for it, val in enumerate(occupancies)
            if 0 < val < 2
        ]
    # HOMO and LUMO plus specified number of orbitals
    # either side
    elif uargs.homo_lumo is not None:
        keep = np.concatenate([
            np.arange(
                np.argmin(occupancies) - 1 - uargs.homo_lumo,
                np.argmin(occupancies),
                1,
                dtype=int
            ),
            np.arange(
                np.argmin(occupancies),
                np.argmin(occupancies) + uargs.homo_lumo + 1,
                1,
                dtype=int
            )
        ])
    elif uargs.num is not None:
        keep = uargs.num
    else:
        keep = range(len(contributions))

    # Remove orbital indices which do not exist
    keep = [val for val in keep if val < len(contributions)]

    if not len(keep):
        ut.red_exit(
            (
                r'Selected orbital indices do not exist!'
                f'\nNORBS = {len(contributions):d}'
            )
        )

    contributions = contributions.loc[:, keep]

    # Remove contributions from unwanted orbitals
    # and shells
    _ORB = {
        's': ['s'],
        'p': ['px', 'py', 'pz'],
        'd': ['dx2y2', 'dz2', 'dxy', 'dxz', 'dyz'],
        'f': ['f0', 'f1', 'f2', 'f3', 'f-1', 'f-2', 'f-3'],
    }
    _orbs_to_use = [
        val
        for chosen in uargs.orb
        for val in _ORB[chosen]
    ]
    if uargs.flavour not in ['redorb_pop', 'orb_comp']:
        _shellorbs_to_use = [
            f'{shell:1d}{orb}'
            for orb in _orbs_to_use
            for shell in uargs.shell
        ]
    else:
        _shellorbs_to_use = _orbs_to_use
    _orb_query = f"AO in {_shellorbs_to_use}"
    contributions = contributions.query(_orb_query)

    # Print info to screen
    ut.cprint(f'Using {extractor_names[uargs.flavour]}', 'cyan')
    if uargs.flavour in ['redorb_pop', 'orb_pop'] and uargs.spin is not None:
        ut.cprint(f'For SPIN={uargs.spin}', 'cyan')
    for mo_num, mo in contributions.items():
        # if no rows greater than threshold, skip
        if not len(mo[mo > uargs.threshold]):
            continue
        _output = ''
        total = 0.
        for row, val in mo.items():
            if val > uargs.threshold and row[1] in uargs.elements:
                _output += f'  {row[1]+str(row[0]):5} {row[2]:5} : {val:>5.1f} %\n' # noqa
                total += val
        if len(_output):
            print(f'MO #{mo_num} (Occ={occupancies[mo_num]}, E={energies[mo_num]: .5f}):') # noqa
            print(_output)
            print(f'  Total:        {total:>5.1f} %\n')

    return


def parse_cutoffs(cutoffs):

    if len(cutoffs) % 2:
        raise argparse.ArgumentTypeError('Error, cutoffs should come in pairs')

    for it in range(1, len(cutoffs), 2):
        try:
            float(cutoffs[it])
        except ValueError:
            raise argparse.ArgumentTypeError(
                'Error, second part of cutoff pair should be float'
            )

    parsed = {}

    for it in range(0, len(cutoffs), 2):
        parsed[cutoffs[it].capitalize()] = float(cutoffs[it + 1])

    return parsed


def extract_freq_func(uargs, save=True):
    '''
    Wrapper for command line frequency extract/print

    Parameters
    ----------
    args : argparser object
        command line arguments
    save : bool, default=True
        If True, saves data to file. If False, prints to stdout.

    Returns
    -------
        None
    '''
    from . import extractor as oe

    # Extract frequency information
    data = oe.FrequencyExtractor.extract(uargs.output_file)

    if not len(data):
        ut.red_exit(f'Cannot find frequencies in file {uargs.output_file}')

    if uargs.num is None:
        uargs.num = len(data[0]['energy (cm^-1)'])

    if not save:
        print('Frequencies (cm⁻¹) and intensities (km/mol)')
        for frq, inty in zip(
            data[0]['energy (cm^-1)'][:uargs.num],
            data[0]['IR Intensity (km mol^-1)'][:uargs.num]
        ):
            print(f'{frq:.5f} {inty:.5f}')
    else:
        # Save to new .csv file
        out_name = f'{uargs.output_file.stem}_frequencies.csv'
        with open(out_name, 'w') as f:
            f.write(
                f'# Frequencies and intensities from {uargs.output_file}\n')
            f.write('# Frequency (cm⁻¹), Intensity (km/mol)\n')
            for frq, inty in zip(
                data[0]['energy (cm^-1)'][:uargs.num],
                data[0]['IR Intensity (km mol^-1)'][:uargs.num]
            ):
                f.write(f'{frq:.5f}, {inty:.5f}\n')

        ut.cprint(f'Data written to {out_name}', 'cyan')

    return


def extract_pop_func(uargs, save=True) -> None:
    '''
    Wrapper for command line frequency extract/print

    Parameters
    ----------
    uargs : argparser object
        User arguments
    save : bool, default=False
        If True, saves data to file. If False, prints to stdout.
    Returns
    -------
        None
    '''
    from . import extractor as oe

    if uargs.flavour in ['loewdin', 'lowdin']:
        data = oe.LoewdinPopulationExtractor.extract(
            uargs.output_file
        )
    elif uargs.flavour in ['mulliken']:
        try:
            data = oe.MullikenPopulationExtractorDensities.extract(
                uargs.output_file
            )
        except DataNotFoundError:
            data = oe.MullikenPopulationExtractorPopulations.extract(
                uargs.output_file
            )

    # Extract structure
    labels, coords = oe.get_coords(uargs.output_file)

    labels = xyzp.add_label_indices(
        labels,
        start_index=0,
        style='sequential'
    )

    if uargs.cutoffs:
        cutoffs = parse_cutoffs(uargs.cutoffs)
    else:
        cutoffs = {}

    # Generate dictionary of entities
    entities_dict = xyzp.find_entities(
        labels, coords, adjust_cutoff=cutoffs, non_bond_labels=uargs.no_bond
    )

    if len(data) > 1:
        ut.cprint(f'Found {len(data)} population blocks in file...', 'green')

    for it, datum in enumerate(data):

        if len(data) > 1:
            ut.cprint(f'\nBlock {it + 1:d}/{len(data):d}', 'green')

        # Calculate charge and spin density of each fragment
        ut.cprint(f'{uargs.flavour.capitalize()} Population Analysis', 'cyan')
        ut.cprint('Entity: Charge Spin', 'cyan')
        ut.cprint('-------------------', 'cyan')
        print()
        for entity_name, entities in entities_dict.items():
            for entity in entities:
                _chg = sum([datum[0][labels[ind]] for ind in entity])
                _spin = sum([datum[1][labels[ind]] for ind in entity])
                ut.cprint(
                    f'{entity_name} : {_chg:.4f}  {_spin:.4f}',
                    'cyan'
                )

    return


def plot_susc_func(uargs) -> None:
    '''
    Plots susceptibility data from output file

    Parameters
    ----------
    args : argparser object
        command line arguments

    Returns
    -------
        None

    '''
    import matplotlib.pyplot as plt
    import matplotlib as mpl
    from . import plotter
    from . import extractor as oe

    # Change matplotlib font size to be larger
    mpl.rcParams.update({'font.size': 12})

    # Extract data from file
    data = oe.SusceptibilityExtractor.extract(uargs.output_file)

    if not len(data):
        ut.red_exit(
            f'Cannot find susceptibility output in {uargs.output_file}'
        )

    if not uargs.quiet:
        ut.cprint(
            f'Found {len(data)} susceptibility blocks in file...',
            'green'
        )
        ut.cprint(
            '... plotting each separately',
            'green'
        )

    if uargs.exp_file is not None:
        exp_data = {'Temperature (K)': [], 'chi*T (cm3*K/mol)': []}
        with open(uargs.exp_file, newline='') as csvfile:
            reader = csv.DictReader(
                row for row in csvfile if not row.startswith('#')
            )
            for row in reader:
                exp_data['Temperature (K)'].append(
                    float(row['Temperature (K)'])
                )
                exp_data['chi*T (cm3*K/mol)'].append(
                    float(row['chi*T (cm3*K/mol)'])
                )

    # Conversion factors from cm3 K mol^-1 to ...
    convs = {
        'A3 K': 1E24 / cst.AVOGADRO,
        'A3 mol-1 K': 1E24,
        'cm3 K': 1 / cst.AVOGADRO,
        'cm3 mol-1 K': 1,
        'emu K': 1 / (4 * np.pi * cst.AVOGADRO),
        'emu mol-1 K': 1 / (4 * np.pi)
    }

    unit_labels = {
        'A3 K': r'\AA^3\,K',
        'A3 mol-1 K': r'\AA^3\,mol^{-1}\,K',
        'cm3 K': r'cm^3\,K',
        'cm3 mol-1 K': r'cm^3\,mol^{-1}\,K',
        'emu K': r'emu\,K',
        'emu mol-1 K': r'emu\,mol^{-1} \ K',
    }

    for dataframe in data:

        fig, ax = plotter.plot_chit(
            dataframe['chi*T (cm3*K/mol)'] * convs[uargs.susc_units],
            dataframe['Temperature (K)'],
            fields=dataframe['Static Field (Gauss)'],
            window_title=f'Susceptibility from {uargs.output_file}',
            y_unit=unit_labels[uargs.susc_units],
            show=_SHOW_CONV[uargs.plot] if uargs.exp_file is None else False,
            save=_SAVE_CONV[uargs.plot] if uargs.exp_file is None else False,
        )
        if uargs.exp_file is not None:
            ax.plot(
                exp_data['Temperature (K)'],
                [
                    val * convs[uargs.esusc_units]
                    for val in exp_data['chi*T (cm3*K/mol)']
                ],
                lw=0,
                marker='o',
                fillstyle='none',
                color='k',
                label='Experiment'
            )
            fig.tight_layout()
            ax.legend(frameon=False)

            _ylim = ax.get_ylim()
            ax.set_ylim(0, _ylim[1])

            if _SAVE_CONV[uargs.plot]:
                plt.savefig('chit_vs_t.png', dpi=500)

            if _SHOW_CONV[uargs.plot]:
                plt.show()

    return


def plot_ailft_func(uargs) -> None:
    '''
    Plots AI-LFT orbital energies

    Parameters
    ----------
    args : argparser object
        command line arguments

    Returns
    -------
        None
    '''
    import matplotlib.pyplot as plt
    import matplotlib as mpl
    from . import plotter
    from . import extractor as oe

    # Change matplotlib font size to be larger
    mpl.rcParams.update({'font.size': 12})

    # Create extractor
    data = oe.AILFTOrbEnergyExtractor.extract(uargs.output_file)

    # Conversion factors from cm-1 to ...
    convs = {
        'cm-1': 1,
        'K': 1E24,
    }

    unit_labels = {
        'cm-1': r'cm^{-1}',
        'K': r'K',
    }

    for dit, dataframe in enumerate(data):

        if len(data) > 1:
            print()
            print(f'Section {dit+1:d}')
            print('---------')

        wfuncs = 100 * np.abs(dataframe['eigenvectors']) ** 2

        for e, wf in zip(dataframe['energies (cm^-1)'], wfuncs.T):
            print(f'E = {e * convs[uargs.units]} {uargs.units}:')
            for it, pc in enumerate(wf):
                if pc > 5.:
                    print(f'{pc:.1f} % {dataframe['orbitals'][it]}')
            if dit == len(wf):
                print('******')
            else:
                print()

        # mm_orbnames = ut.orbname_to_mathmode(dataframe['orbitals'])
        plotter.plot_ailft_orb_energies(
            dataframe['energies (cm^-1)'] * convs[uargs.units],
            groups=uargs.groups,
            occupations=uargs.occupancies,
            # labels=mm_orbnames, # convert these to %ages
            window_title=f'AI-LFT orbitals from {uargs.output_file}',
            y_unit=unit_labels[uargs.units],
            show=_SHOW_CONV[uargs.plot],
            save=_SAVE_CONV[uargs.plot],
            save_name=f'{uargs.output_file.stem}_ailft_orbs_set_{dit+1:d}.png'
        )

        plt.show()

    return


class CustomErrorArgumentParser(argparse.ArgumentParser):
    '''
    Custom ArgumentParser to handle errors and print usage\n
    This is required to avoid the default behavior of argparse which
    modifies the usage message when it prints, conflicting with the preset
    values used in the subparsers.
    '''
    def error(self, message):
        self.print_usage(sys.stderr)
        sys.stderr.write(f"error: {message}.\n")
        sys.stderr.write("       Use -h to see all options.\n")
        sys.exit(2)


def read_args(arg_list=None):
    '''
    Reader for command line arguments. Uses subReaders for individual programs

    Parameters
    ----------
    args : argparser object
        command line arguments

    Returns
    -------
        None
    '''

    description = 'OrcaTools (orto) - A package for working with Orca'

    epilog = 'Type\n'
    epilog += ut.cstring('orto <subprogram> -h\n', 'cyan')
    epilog += 'for help with a specific subprogram.\n'

    parser = CustomErrorArgumentParser(
        usage=ut.cstring('orto <subprogram> [options]', 'cyan'),
        description=description,
        epilog=epilog,
        formatter_class=argparse.RawTextHelpFormatter
    )

    parser._positionals.title = 'Subprograms'

    all_subparsers = parser.add_subparsers(dest='prog_grp')

    extract_subprog = all_subparsers.add_parser(
        'extract',
        description='Extract information from Orca file(s)',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto extract <section>', 'cyan'),
    )

    extract_parser = extract_subprog.add_subparsers(dest='extract_grp')

    extract_subprog._positionals.title = 'Sections'

    # If argument list is empty then call help function
    extract_subprog.set_defaults(func=lambda _: extract_subprog.print_help())

    extract_coords = extract_parser.add_parser(
        'coords',
        description='Extracts coordinates from Orca output file',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto extract coords <output_file> [options]', 'cyan')
    )
    extract_coords._positionals.title = 'Mandatory Arguments'
    extract_coords.set_defaults(func=extract_coords_func)

    extract_coords.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name'
    )

    extract_coords.add_argument(
        '--type',
        type=str,
        help='Which coordinates to extract',
        choices=['opt', 'init'],
        default='init'
    )

    extract_coords.add_argument(
        '--index_style',
        type=str,
        help='Style of indexing used for output atom labels',
        choices=['per_element', 'sequential', 'sequential_orca', 'none'],
        default='per_element'
    )

    extract_freq = extract_parser.add_parser(
        'freq',
        description='Extracts frequencies',
        usage=ut.cstring('orto extract freq <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )

    extract_freq._positionals.title = 'Mandatory Arguments'
    extract_freq.set_defaults(func=extract_freq_func)

    extract_freq.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name - must contain Frequencies section'
    )

    extract_freq.add_argument(
        '-n',
        '--num',
        type=int,
        default=None,
        help='Number of frequencies to print, default is all'
    )

    extract_gmatrix = extract_parser.add_parser(
        'gmatrix',
        description='Extracts coordinates from Orca output file',
        usage=ut.cstring('orto extract gmatrix <output_file> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    extract_gmatrix._positionals.title = 'Mandatory Arguments'
    extract_gmatrix.set_defaults(func=extract_gmatrix_func)

    extract_gmatrix.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name containing G-MATRIX block'
    )

    extract_gmatrix.add_argument(
        '--type',
        type=str,
        help='Which G-MATRIX block to extract.',
        choices=['total', 'S', 'L', 'eff'],
        default='total'
    )

    extract_gmatrix.add_argument(
        '--output_format',
        type=str,
        help='Format of outputted data file',
        choices=['txt', 'docx'],
        default='txt'
    )

    extract_sfenergies = extract_parser.add_parser(
        'sf_energies',
        description='Extract Spin-Free energies from Orca output file',
        usage=ut.cstring('orto extract sf_energies <output_file> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    extract_sfenergies._positionals.title = 'Mandatory Arguments'
    extract_sfenergies.set_defaults(func=extract_sf_energies_func)

    extract_sfenergies.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name containing TRANSITION ENERGIES block'
    )

    extract_sfenergies.add_argument(
        '--output_format',
        type=str,
        help='Format of outputted data file',
        choices=['txt', 'docx'],
        default='txt'
    )

    extract_soenergies = extract_parser.add_parser(
        'so_energies',
        description='Extract Spin-Free energies from Orca output file',
        usage=ut.cstring('orto extract so_energies <output_file> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    extract_soenergies._positionals.title = 'Mandatory Arguments'
    extract_soenergies.set_defaults(func=extract_so_energies_func)

    extract_soenergies.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name containing Spin Orbit rel block'
    )

    gen_subprog = all_subparsers.add_parser(
        'gen',
        description='Generate inputs/jobs to/for Orca',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto gen <item>', 'cyan')
    )
    gen_subprog._positionals.title = 'Items'

    gen_parser = gen_subprog.add_subparsers(dest='gen_grp')

    # If argument list is empty then call help function
    gen_subprog.set_defaults(func=lambda _: gen_subprog.print_help())

    gen_trunc_molden = gen_parser.add_parser(
        'trunc_molden',
        description='Generate truncated molden file',
        usage=ut.cstring('orto gen trunc_molden <input_file> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    gen_trunc_molden._positionals.title = 'Mandatory Arguments'
    gen_trunc_molden.set_defaults(func=gen_trunc_molden_func)
    gen_trunc_molden.add_argument(
        'input_file',
        type=pathlib.Path,
        help='Name of molden file to truncate'
    )

    gen_trunc_molden.add_argument(
        '--output_file',
        type=pathlib.Path,
        help=(
            'Name of truncated molden file\n'
            'If not specified, molden file will be truncated in place'
        ),
        default=None
    )

    gen_trunc_molden.add_argument(
        '--n_virt',
        type=int,
        default=100,
        help=(
            'Number of virtual orbitals to keep in truncated molden file\n'
            'Default: %(default)s'
        )
    )

    gen_job = gen_parser.add_parser(
        'job',
        description=(
            'Generate submission script for orca calculation.\n'
            'Job script should be executed/submitted from its parent directory'
        ),
        usage=ut.cstring('orto gen job <input_file(s)> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    gen_job._positionals.title = 'Mandatory Arguments'

    gen_job.set_defaults(func=gen_job_func)

    gen_job.add_argument(
        'input_files',
        metavar='<input_file(s)>',
        type=pathlib.Path,
        nargs='+',
        help='Orca input file name(s)'
    )

    gen_job.add_argument(
        '--n_procs',
        type=int,
        default=0,
        help=(
            'Number of cores requested in submission system.\n'
            ' This does not need to match the orca input, but must not be'
            'less\n. If not specified then value is read from input file.')
    )

    gen_job.add_argument(
        '--memory',
        '-mem',
        type=int,
        default=0,
        help=(
            'Per core memory requested in submission system (megabytes).\n'
            ' This does not need to match the orca input, but must not be'
            ' less.\n If not specified then value is read from input file.'
        )
    )

    gen_job.add_argument(
        '--no_sub',
        '-ns',
        action='store_true',
        help=(
            'Disables submission of job to queue'
        )
    )

    gen_job.add_argument(
        '--no_molden',
        '-nm',
        action='store_true',
        help=(
            'Disables orca_2mkl call for molden file generation after calculation' # noqa
        )
    )

    gen_job.add_argument(
        '--skip_xyz',
        '-sx',
        action='store_true',
        help=(
            'Disables xyz file format check'
        )
    )

    gen_job.add_argument(
        '-om',
        '--orca_load',
        type=str,
        default='',
        help='Orca environment module (overrides ORTO_ORCA_LOAD envvar)'
    )

    distort = gen_parser.add_parser(
        'distort',
        description='Distorts molecule along given normal mode',
        usage=ut.cstring('orto gen distort <output_file> <mode_number> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    distort._positionals.title = 'Mandatory Arguments'

    distort.set_defaults(func=distort_func)

    distort.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name - must contain frequency section'
    )

    distort.add_argument(
        'mode_number',
        type=int,
        help='Mode to distort along - uses orca indexing and starts from zero'
    )

    distort.add_argument(
        '--scale',
        type=float,
        default=1,
        help=(
            'Number of units of distortion\n'
            'Default: %(default)s'
        )
    )

    plot_subprog = all_subparsers.add_parser(
        'plot',
        description='Plot data from orca file',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto plot <section>', 'cyan'),
    )
    plot_subprog._positionals.title = 'Mandatory Arguments'

    plot_parser = plot_subprog.add_subparsers(dest='plot_grp')

    # If argument list is empty then call help function
    plot_subprog.set_defaults(func=lambda _: plot_subprog.print_help())

    plot_abs = plot_parser.add_parser(
        'abs',
        description='Plots absorption spectrum from CI calculation output',
        usage=ut.cstring('orto plot abs <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    plot_abs._positionals.title = 'Mandatory Arguments'

    plot_abs.set_defaults(func=plot_abs_func)

    plot_abs.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name'
    )

    plot_abs.add_argument(
        '--intensities',
        '-i',
        type=str,
        choices=['velocity', 'electric', 'semi-classical'],
        default='electric',
        help='Type of intensity to plot (orca_mapspc uses electric)'
    )

    plot_abs.add_argument(
        '--linewidth',
        '-lw',
        type=float,
        default=2000,
        help=(
            'Width of signal (FWHM for Gaussian, Width for Lorentzian),'
            ' in Wavenumbers'
        )
    )

    plot_abs.add_argument(
        '--osc_style',
        type=str,
        default='combined',
        help=(
            'Style of oscillators to plot\n'
            ' - \'separate\' plots oscillator strengths as stems on separate axis\n'
            ' - \'combined\' plots oscillator strengths on intensity axis\n'
            ' - \'off\' does not plot oscillator strengths\n'
            'Default: %(default)s'
        )
    )

    plot_abs.add_argument(
        '--plot',
        '-p',
        choices=['on', 'show', 'save', 'off'],
        metavar='<str>',
        type=str,
        default='on',
        help=(
            'Controls plot appearance/save \n'
            ' - \'on\' shows and saves the plots\n'
            ' - \'show\' shows the plots\n'
            ' - \'save\' saves the plots\n'
            ' - \'off\' neither shows nor saves\n'
            'Default: %(default)s'
        )
    )

    plot_abs.add_argument(
        '--lineshape',
        '-ls',
        type=str,
        choices=['gaussian', 'lorentzian'],
        default='lorentzian',
        help='Lineshape to use for each signal'
    )

    plot_abs.add_argument(
        '--x_unit',
        type=str,
        choices=['wavenumber', 'energy', 'wavelength'],
        default='wavenumber',
        help='x units to use for spectrum'
    )

    plot_abs.add_argument(
        '--shift',
        type=float,
        default=0.,
        help=(
            'Shift spectrum by this amount in x units\n'
            'Default: %(default)s'
        )
    )

    plot_abs.add_argument(
        '--x_lim',
        nargs=2,
        default=['auto', 'auto'],
        help='x limits of spectrum'
    )

    plot_abs.add_argument(
        '--y_lim',
        nargs=2,
        default=[0., 'auto'],
        help='Epsilon limits of spectrum in cm^-1 mol^-1 L'
    )

    plot_abs.add_argument(
        '--normalise_absorption',
        '-na',
        action='store_true',
        default=False,
        help=(
            'Normalises absorption spectrum to maximum value\n'
            'Default: %(default)s'
        )
    )

    plot_ailft = plot_parser.add_parser(
        'ailft_orbs',
        description='Plots AI-LFT orbital energies from output file',
        usage=ut.cstring(
            'orto plot ailft_orbs <output_file> [options]',
            'cyan'
        ),
        formatter_class=argparse.RawTextHelpFormatter
    )
    plot_ailft._positionals.title = 'Mandatory Arguments'

    plot_ailft.set_defaults(func=plot_ailft_func)

    plot_ailft.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name'
    )

    plot_ailft.add_argument(
        '--groups',
        '-g',
        metavar='<str>',
        nargs='+',
        type=int,
        default=None,
        help=(
            'Group indices for each orbital. e.g. 1 1 1 2 2\n'
            'Controls x-staggering of orbitals'
        )
    )

    plot_ailft.add_argument(
        '--occupancies',
        '-o',
        metavar='<str>',
        nargs='+',
        type=int,
        default=None,
        help=(
            'Occupation number of each orbital\n Adds electrons to each orb\n'
            'Must specify occupation of every orbital as 2, 1, -1, or 0'
        )
    )

    plot_ailft.add_argument(
        '--plot',
        '-p',
        choices=['on', 'show', 'save', 'off'],
        metavar='<str>',
        type=str,
        default='on',
        help=(
            'Controls plot appearance/save \n'
            ' - \'on\' shows and saves the plots\n'
            ' - \'show\' shows the plots\n'
            ' - \'save\' saves the plots\n'
            ' - \'off\' neither shows nor saves\n'
            'Default: %(default)s'
        )
    )

    plot_ailft.add_argument(
        '--units',
        '-u',
        choices=[
            'cm-1',
            'K'
        ],
        metavar='<str>',
        type=str,
        default='cm-1',
        help=(
            'Controls energy units of plot\n'
            'Default: %(default)s'
        )
    )

    plot_susc = plot_parser.add_parser(
        'susc',
        description='Plots susceptibility data from output file',
        usage=ut.cstring('orto plot susc <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    plot_susc._positionals.title = 'Mandatory Arguments'

    plot_susc.set_defaults(func=plot_susc_func)

    plot_susc.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name'
    )

    plot_susc.add_argument(
        '--susc_units',
        '-su',
        choices=[
            'emu mol-1 K',
            'emu K',
            'cm3 mol-1 K',
            'cm3 K',
            'A3 mol-1 K',
            'A3 K'
        ],
        metavar='<str>',
        type=str,
        default='cm3 mol-1 K',
        help=(
            'Controls susceptibility units of calculated data \n'
            '(wrap with "")\n'
            'Default: %(default)s'
        )
    )

    plot_susc.add_argument(
        '--plot',
        '-p',
        choices=['on', 'show', 'save', 'off'],
        metavar='<str>',
        type=str,
        default='on',
        help=(
            'Controls plot appearance/save \n'
            ' - \'on\' shows and saves the plots\n'
            ' - \'show\' shows the plots\n'
            ' - \'save\' saves the plots\n'
            ' - \'off\' neither shows nor saves\n'
            'Default: %(default)s'
        )
    )

    plot_susc.add_argument(
        '--exp_file',
        type=str,
        help=(
            'Experimental datafile as .csv with two columns:\n'
            '1. "Temperature (K)"\n'
            '2. "chi*T (UNITS FROM --esusc_units)"\n'
        )
    )

    plot_susc.add_argument(
        '--esusc_units',
        '-esu',
        choices=[
            'emu mol-1 K',
            'emu K',
            'cm3 mol-1 K',
            'cm3 K',
            'A3 mol-1 K',
            'A3 K'
        ],
        metavar='<str>',
        type=str,
        default='cm3 mol-1 K',
        help=(
            'Controls susceptibility units of experimental data \n'
            '(wrap with "")\n'
            'Default: %(default)s'
        )
    )

    plot_susc.add_argument(
        '--quiet',
        action='store_true',
        help='Suppresses text output'
    )

    plot_ir = plot_parser.add_parser(
        'ir',
        description='Plots infrared spectrum from frequency calculation output', # noqa
        usage=ut.cstring('orto plot ir <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    plot_ir._positionals.title = 'Mandatory Arguments'

    plot_ir.set_defaults(func=plot_ir_func)

    plot_ir.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name'
    )

    plot_ir.add_argument(
        '--plot',
        '-p',
        choices=['on', 'show', 'save', 'off'],
        metavar='<str>',
        type=str,
        default='on',
        help=(
            'Controls plot appearance/save \n'
            ' - \'on\' shows and saves the plots\n'
            ' - \'show\' shows the plots\n'
            ' - \'save\' saves the plots\n'
            ' - \'off\' neither shows nor saves\n'
            'Default: %(default)s'
        )
    )

    plot_ir.add_argument(
        '--linewidth',
        '-lw',
        type=float,
        default=5,
        help=(
            'Width of signal (FWHM for Gaussian, Width for Lorentzian),'
            ' in same unit as plot x unit'
        )
    )

    plot_ir.add_argument(
        '--lineshape',
        '-ls',
        type=str,
        choices=['gaussian', 'lorentzian'],
        default='lorentzian',
        help='Lineshape to use for each signal'
    )

    print_subprog = all_subparsers.add_parser(
        'print',
        description='Print information from Orca file to screen',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto print <section>', 'cyan')
    )
    print_subprog._positionals.title = 'Sections'

    print_parser = print_subprog.add_subparsers(dest='print_grp')

    # If argument list is empty then call help function
    print_subprog.set_defaults(func=lambda _: print_subprog.print_help())

    print_gmatrix = print_parser.add_parser(
        'gmatrix',
        description='Extracts g matrix from Orca output file',
        usage=ut.cstring('orto print gmatrix <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    print_gmatrix._positionals.title = 'Mandatory Arguments'

    print_gmatrix.set_defaults(
        func=lambda x: extract_gmatrix_func(x, save=False)
    )

    print_gmatrix.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name containing G-MATRIX block'
    )

    print_gmatrix.add_argument(
        '--type',
        type=str,
        help='Which G-MATRIX block to extract.',
        choices=['total', 'S', 'L', 'eff'],
        default='total'
    )

    print_freq = print_parser.add_parser(
        'freq',
        description='Prints frequencies',
        usage=ut.cstring('orto print freq <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    print_freq._positionals.title = 'Mandatory Arguments'

    print_freq.set_defaults(func=lambda x: extract_freq_func(x, save=False))

    print_freq.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name - must contain Frequencies section'
    )

    print_freq.add_argument(
        '-n',
        '--num',
        type=int,
        default=None,
        help='Number of frequencies to print, default is all'
    )

    print_orbs = print_parser.add_parser(
        'lorbs',
        description='Prints Loewdin orbital compositions',
        usage=ut.cstring('orto print orbs <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    print_orbs._positionals.title = 'Mandatory Arguments'

    print_orbs.set_defaults(
        func=lambda x: extract_orbs_func(x, save=False)
    )

    print_orbs.add_argument(
        'output_file',
        type=pathlib.Path,
        help=(
            'Orca output file name\n'
            'File must contain one of the following sections\n'
            '   LOEWDIN ORBITAL-COMPOSITIONS\n'
            '   LOEWDIN REDUCED ORBITAL POPULATIONS PER MO\n'
            '   LOEWDIN ORBITAL POPULATIONS PER MO\n'
        )
    )

    print_orbs.add_argument(
        '-t',
        '--threshold',
        type=float,
        default=1.,
        help=(
            'Orbitals with contribution >= threshold are printed.\n'
            'Default: %(default)s'
        )
    )

    print_orbs.add_argument(
        '-e',
        '--elements',
        type=str,
        default=atomic_elements,
        nargs='+',
        help='Only print contributions from specified element(s) e.g. Ni'
    )

    print_orbs.add_argument(
        '-s',
        '--shell',
        type=int,
        nargs='+',
        default=[1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 19, 20], # noqa
        help=(
            'Only print contributions from specified shell numbers e.g. 3\n'
            'Default: %(default)s'
        )
    )

    print_orbs.add_argument(
        '-o',
        '--orb',
        type=str,
        default=['s', 'p', 'd', 'f'],
        choices=['s', 'p', 'd', 'f'],
        nargs='+',
        help=(
            'Only print contributions from specified orbital(s) e.g. d\n'
            'Default: %(default)s'
        )
    )

    print_orbs.add_argument(
        '-f',
        '--flavour',
        type=str,
        choices=['first_match', 'orb_pop', 'redorb_pop', 'orb_comp'],
        default='first_match',
        help=(
            'Which section to print from the output file\n'
            'orb_comp: Loewdin orbital compositions\n'
            'orb_pop: Loewdin orbital populations\n'
            'redorb_pop: Loewdin reduced orbital populations\n'
        )
    )

    print_orbs.add_argument(
        '--spin',
        type=str,
        choices=['UP', 'DOWN'],
        default='UP',
        help=(
            'Which spin to print from the output file\n'
            'If closed shell, then this is ignored\n'
            'Default: %(default)s'
        )
    )

    orb_group = print_orbs.add_mutually_exclusive_group(required=False)

    orb_group.add_argument(
        '-a',
        '--active',
        action='store_true',
        help=(
            'Only print active orbitals (0 < occupation < 2)'
        )
    )

    def gte_zero(x):
        '''
        Custom type for argparse to ensure that the input
        \nis greater than or equal to zero
        '''
        value = int(x)
        if value < 0:
            raise argparse.ArgumentTypeError(
                f'{x} is not a valid index (must be >= 0)'
            )
        return value

    orb_group.add_argument(
        '-n',
        '--num',
        nargs='+',
        type=gte_zero,
        metavar='NUMBER',
        default=None,
        help=(
            'Print specified orbitals using index starting from 0\n'
            '(same as Orca)\n'
        )
    )

    orb_group.add_argument(
        '-hl',
        '--homo_lumo',
        nargs='?',
        type=int,
        metavar='NUMBER',
        default=None,
        help=(
            'Print specified number of orbitals either side of HOMO and LUMO'
        )
    )

    print_pop = print_parser.add_parser(
        'pop',
        description='Prints population analysis (spin, charge), and groups by fragment', # noqa
        usage=ut.cstring('orto print pop <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    print_pop._positionals.title = 'Mandatory Arguments'

    print_pop.set_defaults(func=lambda x: extract_pop_func(x, save=False))

    print_pop.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name - must contain population analysis section'
    )

    print_pop.add_argument(
        '--flavour',
        '-f',
        type=str,
        choices=['lowdin', 'loewdin', 'mulliken'],
        default='mulliken',
        help='Type of population analysis to print'
    )

    print_pop.add_argument(
        '--no_bond',
        '-nb',
        type=str,
        default=[],
        nargs='+',
        metavar='symbol',
        help='Atom labels specifying atoms to which no bonds can be formed'
    )

    print_pop.add_argument(
        '--cutoffs',
        type=str,
        nargs='+',
        metavar='symbol number',
        help='Modify cutoff used to define bonds between atoms'
    )

    # If argument list is empty then call help function
    parser.set_defaults(func=lambda _: parser.print_help())

    # select parsing option based on sub-parser
    args = parser.parse_args(arg_list)
    args.func(args)
    return args


def main():
    read_args()
    return
