"""
Text extraction utilities for different file formats
"""

import os
import re
from pathlib import Path
from typing import Union

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

class TextExtractorError(Exception):
    """Custom exception for text extraction errors"""
    pass

class TextExtractor:
    """Text extraction utility class"""
    
    @staticmethod
    def extract_text(file_path: Union[str, Path]) -> str:
        """Extract text from various file formats"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise TextExtractorError(f"File not found: {file_path}")
        
        if file_path.stat().st_size == 0:
            raise TextExtractorError(f"File is empty: {file_path}")
        
        ext = file_path.suffix.lower()
        
        if ext == ".pdf":
            return TextExtractor._extract_pdf_text(file_path)
        elif ext == ".docx":
            return TextExtractor._extract_docx_text(file_path)
        elif ext == ".txt":
            return TextExtractor._extract_txt_text(file_path)
        else:
            raise TextExtractorError(f"Unsupported file type: {ext}")
    
    @staticmethod
    def _extract_pdf_text(file_path: Path) -> str:
        """Extract text from PDF"""
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(str(file_path))
                text_blocks = []
                
                for page_num, page in enumerate(doc):
                    page_text = page.get_text()
                    if page_text.strip():
                        text_blocks.append(f"--- Page {page_num + 1} ---\n{page_text}")
                
                doc.close()
                
                if not text_blocks:
                    raise TextExtractorError("No readable text found in PDF")
                
                return "\n\n".join(text_blocks)
            except Exception as e:
                raise TextExtractorError(f"PyMuPDF extraction failed: {e}")
        
        elif PDFMINER_AVAILABLE:
            try:
                text = extract_text(str(file_path))
                if not text or not text.strip():
                    raise TextExtractorError("No readable text found in PDF")
                return text
            except Exception as e:
                raise TextExtractorError(f"PDFMiner extraction failed: {e}")
        
        else:
            raise TextExtractorError("No PDF parsing libraries available. Install PyMuPDF or pdfminer.six")
    
    @staticmethod
    def _extract_docx_text(file_path: Path) -> str:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            raise TextExtractorError("python-docx not available. Install with: pip install python-docx")
        
        try:
            doc = docx.Document(str(file_path))
            text_blocks = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_blocks.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_blocks.append("TABLE:\n" + "\n".join(table_text))
            
            if not text_blocks:
                raise TextExtractorError("No readable text found in DOCX file")
            
            return "\n\n".join(text_blocks)
        
        except Exception as e:
            raise TextExtractorError(f"DOCX extraction failed: {e}")
    
    @staticmethod
    def _extract_txt_text(file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            except UnicodeDecodeError:
                # Fallback to other encodings
                for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                    try:
                        with open(file_path, "r", encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError:
                        continue
                
                raise TextExtractorError("Could not decode text file with any common encoding")
        
        except Exception as e:
            raise TextExtractorError(f"TXT extraction failed: {e}")
