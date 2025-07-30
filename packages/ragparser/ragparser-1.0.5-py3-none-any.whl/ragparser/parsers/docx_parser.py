"""
DOCX parser implementation using python-docx
"""

import asyncio
import logging
from typing import List, Dict, Any, TYPE_CHECKING
from pathlib import Path
import io

try:
    from docx import Document
    from docx.shared import Inches

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    # Create dummy class for type hints
    if TYPE_CHECKING:
        from docx import Document
    else:
        Document = object

from .base import BaseParser
from ..core.models import ParsedDocument, ParserConfig, FileType, ContentBlock
from ..core.exceptions import ProcessingError


class DOCXParser(BaseParser):
    """Parser for DOCX documents"""

    def __init__(self):
        super().__init__()
        self.supported_formats = [FileType.DOCX]

        if not HAS_DOCX:
            raise ImportError(
                "DOCX parsing requires python-docx. Install with: pip install python-docx"
            )

    async def parse_async(
        self, file_path: Path, config: ParserConfig
    ) -> ParsedDocument:
        """Parse DOCX file asynchronously"""
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(
                None, self._parse_docx_sync, str(file_path), config
            )
        except Exception as e:
            raise ProcessingError(f"Failed to parse DOCX: {str(e)}", e)

    async def parse_from_bytes_async(
        self, data: bytes, filename: str, config: ParserConfig
    ) -> ParsedDocument:
        """Parse DOCX from bytes asynchronously"""
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(
                None, self._parse_docx_from_bytes_sync, data, filename, config
            )
        except Exception as e:
            raise ProcessingError(f"Failed to parse DOCX from bytes: {str(e)}", e)

    def _parse_docx_sync(self, file_path: str, config: ParserConfig) -> ParsedDocument:
        """Synchronous DOCX parsing"""
        path = Path(file_path)
        document = asyncio.run(
            self._create_base_document(
                path, path.name, FileType.DOCX, path.stat().st_size
            )
        )

        try:
            doc = Document(file_path)
            return self._extract_content(doc, document, config)
        except Exception as e:
            raise ProcessingError(f"Failed to parse DOCX: {str(e)}", e)

    def _parse_docx_from_bytes_sync(
        self, data: bytes, filename: str, config: ParserConfig
    ) -> ParsedDocument:
        """Synchronous DOCX parsing from bytes"""
        document = asyncio.run(
            self._create_base_document(None, filename, FileType.DOCX, len(data))
        )

        try:
            doc = Document(io.BytesIO(data))
            return self._extract_content(doc, document, config)
        except Exception as e:
            raise ProcessingError(f"Failed to parse DOCX from bytes: {str(e)}", e)

    def _extract_content(
        self, doc, document: ParsedDocument, config: ParserConfig
    ) -> ParsedDocument:
        """Extract content from DOCX document"""
        """Extract content from DOCX document"""
        content_parts = []
        content_blocks = []
        tables = []
        images = []
        links = []

        # Extract core properties
        if doc.core_properties:
            props = doc.core_properties
            document.metadata.title = props.title
            document.metadata.author = props.author
            document.metadata.subject = props.subject
            document.metadata.creation_date = props.created
            document.metadata.modification_date = props.modified

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                block_type = "text"

                # Detect headers based on style
                if para.style.name.startswith("Heading"):
                    block_type = "header"
                elif para.style.name == "Title":
                    block_type = "title"

                content_block = ContentBlock(
                    content=text,
                    block_type=block_type,
                    formatting={
                        "style": para.style.name,
                        "bold": any(run.bold for run in para.runs),
                        "italic": any(run.italic for run in para.runs),
                        "underline": any(run.underline for run in para.runs),
                    },
                )

                content_blocks.append(content_block)
                content_parts.append(text)

                # Extract hyperlinks if enabled
                if config.extract_links:
                    for run in para.runs:
                        if hasattr(run.element, "hyperlink") and run.element.hyperlink:
                            try:
                                rel = doc.part.rels[run.element.hyperlink.get("r:id")]
                                if rel.target_ref:
                                    links.append(
                                        {"text": run.text, "url": rel.target_ref}
                                    )
                            except:
                                pass

        # Extract tables if enabled
        if config.extract_tables:
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        table_data.append(row_data)

                    if table_data:
                        tables.append({"index": table_idx, "data": table_data})

                        # Add table content to text
                        table_text = "\n".join(["\t".join(row) for row in table_data])
                        content_blocks.append(
                            ContentBlock(content=table_text, block_type="table")
                        )
                        content_parts.append(table_text)

                except Exception as e:
                    logging.warning(f"Failed to extract table {table_idx}: {str(e)}")

        # Extract images if enabled
        if config.extract_images:
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        images.append({"name": rel.target_ref, "type": rel.reltype})
            except Exception as e:
                logging.warning(f"Failed to extract images: {str(e)}")

        # Combine content
        if config.merge_paragraphs:
            document.content = "\n\n".join(content_parts)
        else:
            document.content = "\n".join(content_parts)

        document.content_blocks = content_blocks
        document.tables = tables
        document.images = images
        document.links = links

        # Update metadata
        document.metadata.word_count = len(document.content.split())
        document.metadata.character_count = len(document.content)

        return document
