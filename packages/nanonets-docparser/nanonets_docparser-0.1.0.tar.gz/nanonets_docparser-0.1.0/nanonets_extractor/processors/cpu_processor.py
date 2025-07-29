"""
CPU processor using local models for document extraction.
"""

import os
import json
import re
from typing import Dict, List, Any, Optional
from pathlib import Path

# Document processing libraries
try:
    import PyPDF2
    import docx
    import openpyxl
    import pandas as pd
    from PIL import Image
    import pytesseract
    import easyocr
    import cv2
    import numpy as np
except ImportError as e:
    print(f"Warning: Some dependencies not available for CPU processing: {e}")

from .base_processor import BaseProcessor
from ..utils import OutputType, FileType
from ..exceptions import ModelError, ExtractionError


class CPUProcessor(BaseProcessor):
    """
    Processor that uses local CPU models for document extraction.
    """
    
    def __init__(self, model_path: Optional[str] = None, **kwargs):
        """
        Initialize the CPU processor.
        
        Args:
            model_path: Path to custom models
            **kwargs: Additional configuration
        """
        super().__init__(**kwargs)
        self.model_path = model_path
        self._ocr_reader = None
        self._initialize_models()
    
    def _initialize_models(self):
        """Initialize local models."""
        try:
            # Initialize EasyOCR for text extraction
            self._ocr_reader = easyocr.Reader(['en'])
        except Exception as e:
            print(f"Warning: EasyOCR not available: {e}")
            self._ocr_reader = None
    
    def is_available(self) -> bool:
        """Check if CPU processor is available."""
        return True  # CPU processing is always available
    
    def extract(
        self,
        file_path: str,
        file_type: FileType,
        output_type: OutputType,
        specified_fields: Optional[List[str]] = None,
        json_schema: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Extract data using local CPU processing.
        
        Args:
            file_path: Path to the document file
            file_type: Type of the file
            output_type: Desired output format
            specified_fields: List of fields to extract
            json_schema: Custom JSON schema
            
        Returns:
            Dictionary containing extracted data
        """
        try:
            # Extract text based on file type
            if file_type == FileType.PDF:
                extracted_data = self._extract_from_pdf(file_path)
            elif file_type == FileType.IMAGE:
                extracted_data = self._extract_from_image(file_path)
            elif file_type == FileType.DOCUMENT:
                extracted_data = self._extract_from_document(file_path)
            elif file_type == FileType.SPREADSHEET:
                extracted_data = self._extract_from_spreadsheet(file_path)
            elif file_type == FileType.TEXT:
                extracted_data = self._extract_from_text(file_path)
            else:
                raise ExtractionError(f"Unsupported file type: {file_type}")
            
            # Format output based on output_type
            return self._format_output(extracted_data, output_type, specified_fields, json_schema)
            
        except Exception as e:
            raise ExtractionError(f"CPU extraction failed: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    "text": text,
                    "pages": len(pdf_reader.pages),
                    "file_type": "pdf"
                }
        except Exception as e:
            raise ModelError(f"PDF extraction failed: {e}")
    
    def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image file using OCR."""
        try:
            # Load image
            image = cv2.imread(file_path)
            if image is None:
                raise ModelError(f"Could not load image: {file_path}")
            
            # Convert to RGB for EasyOCR
            image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            
            # Extract text using EasyOCR
            if self._ocr_reader:
                results = self._ocr_reader.readtext(image_rgb)
                text = " ".join([result[1] for result in results])
            else:
                # Fallback to Tesseract
                text = pytesseract.image_to_string(image_rgb)
            
            return {
                "text": text,
                "file_type": "image",
                "ocr_confidence": 0.8  # Placeholder
            }
        except Exception as e:
            raise ModelError(f"Image OCR failed: {e}")
    
    def _extract_from_document(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Word document."""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                "text": text,
                "file_type": "document",
                "paragraphs": len(doc.paragraphs)
            }
        except Exception as e:
            raise ModelError(f"Document extraction failed: {e}")
    
    def _extract_from_spreadsheet(self, file_path: str) -> Dict[str, Any]:
        """Extract data from spreadsheet."""
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Convert to structured data
            data = {
                "columns": df.columns.tolist(),
                "rows": df.values.tolist(),
                "shape": df.shape,
                "file_type": "spreadsheet"
            }
            
            # Also extract as text for compatibility
            data["text"] = df.to_string()
            
            return data
        except Exception as e:
            raise ModelError(f"Spreadsheet extraction failed: {e}")
    
    def _extract_from_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                "text": text,
                "file_type": "text",
                "length": len(text)
            }
        except Exception as e:
            raise ModelError(f"Text file extraction failed: {e}")
    
    def _format_output(
        self,
        extracted_data: Dict[str, Any],
        output_type: OutputType,
        specified_fields: Optional[List[str]] = None,
        json_schema: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Format extracted data according to output type."""
        if output_type == OutputType.MARKDOWN:
            return self._format_markdown(extracted_data)
        elif output_type == OutputType.FLAT_JSON:
            return self._format_flat_json(extracted_data)
        elif output_type == OutputType.SPECIFIED_FIELDS:
            return self._format_specified_fields(extracted_data, specified_fields)
        elif output_type == OutputType.SPECIFIED_JSON:
            return self._format_specified_json(extracted_data, json_schema)
        else:
            return extracted_data
    
    def _format_markdown(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Format data as markdown."""
        text = data.get("text", "")
        return {
            "markdown": text,
            "file_type": data.get("file_type", "unknown")
        }
    
    def _format_flat_json(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Format data as flat JSON with extracted fields."""
        result = {}
        
        # Extract common fields from text
        text = data.get("text", "")
        
        # Simple field extraction patterns
        patterns = {
            "invoice_number": r"invoice[:\s#]*(\d+)",
            "date": r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            "amount": r"(\$?\d+\.?\d*)",
            "email": r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})",
            "phone": r"(\(\d{3}\)\s?\d{3}-\d{4}|\d{3}-\d{3}-\d{4})",
        }
        
        for field, pattern in patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                result[field] = matches[0] if len(matches) == 1 else matches
        
        # Add basic metadata
        result.update({
            "file_type": data.get("file_type", "unknown"),
            "extracted_text": text[:1000] + "..." if len(text) > 1000 else text,
        })
        
        return result
    
    def _format_specified_fields(
        self, 
        data: Dict[str, Any], 
        specified_fields: Optional[List[str]]
    ) -> Dict[str, Any]:
        """Format data with only specified fields."""
        if not specified_fields:
            return {}
        
        flat_data = self._format_flat_json(data)
        result = {}
        
        for field in specified_fields:
            if field in flat_data:
                result[field] = flat_data[field]
            else:
                result[field] = None
        
        return result
    
    def _format_specified_json(
        self, 
        data: Dict[str, Any], 
        json_schema: Optional[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Format data according to custom JSON schema."""
        if not json_schema:
            return {}
        
        flat_data = self._format_flat_json(data)
        result = {}
        
        # Try to map extracted data to schema
        for field, field_type in json_schema.items():
            if field in flat_data:
                value = flat_data[field]
                
                # Type conversion based on schema
                if field_type == "number" and isinstance(value, str):
                    try:
                        # Remove currency symbols and convert
                        clean_value = re.sub(r'[^\d.]', '', value)
                        result[field] = float(clean_value)
                    except:
                        result[field] = 0.0
                elif field_type == "string":
                    result[field] = str(value) if value else ""
                else:
                    result[field] = value
            else:
                # Set default values based on type
                if field_type == "number":
                    result[field] = 0.0
                elif field_type == "string":
                    result[field] = ""
                elif field_type == "array":
                    result[field] = []
                else:
                    result[field] = None
        
        return result
    
    def get_capabilities(self) -> Dict[str, Any]:
        """Get CPU processor capabilities."""
        capabilities = super().get_capabilities()
        capabilities.update({
            "ocr_available": self._ocr_reader is not None,
            "supports_pdf": True,
            "supports_images": self._ocr_reader is not None,
            "supports_documents": True,
            "supports_spreadsheets": True,
            "supports_text": True,
            "offline": True,
            "requires_internet": False,
        })
        return capabilities 